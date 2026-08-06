import sys
import os
import re
import json
import sqlite3
import hashlib
import secrets
import math
from datetime import datetime
from pathlib import Path
import subprocess
import tempfile
import shutil
from contextlib import contextmanager
import io
import traceback

from flask import Flask, render_template, request, redirect, url_for, send_file, session, flash, jsonify, make_response
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Используем только PyMuPDF (fitz)
import fitz

PDF_LIB = 'fitz'

# Импорт конфигурации
try:
    from config import Config

    print("✅ Конфигурация загружена из config.py")
except ImportError:
    print("⚠️ config.py не найден, используем настройки по умолчанию")


    class Config:
        SECRET_KEY = secrets.token_hex(32)
        UPLOAD_FOLDER = 'uploads'
        MAX_CONTENT_LENGTH = 16 * 1024 * 1024

app = Flask(__name__)
app.secret_key = Config.SECRET_KEY
app.config['UPLOAD_FOLDER'] = Config.UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = Config.MAX_CONTENT_LENGTH

# Создаем папку для загрузок
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Для Render: сохраняем файлы в папке приложения
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Переназначаем пути к файлам в папку приложения
SETTINGS_FILE = os.path.join(BASE_DIR, 'settings.json')
DATABASE_FILE = os.path.join(BASE_DIR, 'users.db')
MATERIALS_DATABASE_FILE = os.path.join(BASE_DIR, 'materials.db')
LAST_FOLDER_FILE = os.path.join(BASE_DIR, 'last_folder.json')
LAST_FILE_FILE = os.path.join(BASE_DIR, 'last_file.json')


# ============================================================
# УНИВЕРСАЛЬНАЯ ФУНКЦИЯ ИЗВЛЕЧЕНИЯ ТЕКСТА ИЗ PDF
# ============================================================

def extract_text_from_pdf(pdf_path):
    """
    Универсальная функция извлечения текста из PDF.
    Работает с PyMuPDF.
    """
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            text += page_text
        doc.close()
        return text
    except Exception as e:
        print(f"❌ Ошибка открытия PDF: {e}")
        traceback.print_exc()
        return None


def extract_data_from_pdf(pdf_path):
    """Извлекает данные из PDF-файла приходного ордера"""
    # Извлекаем текст из PDF
    text = extract_text_from_pdf(pdf_path)
    if text is None:
        return None

    # Разбиваем на строки и очищаем
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    result = {
        'дата_приемки': None,
        'название_поставщика': None,
        'номер_ордера': None,
        'товары': []
    }

    # 1) Номер приходного ордера
    for line in lines:
        if 'ПРИХОДНЫЙ ОРДЕР №' in line or 'ПРИХОДНЫЙ ОРДЕР N' in line:
            order_match = re.search(r'№?\s*(\d+)', line)
            if order_match:
                result['номер_ордера'] = order_match.group(1)
                break

    # 2) Дата приемки
    for line in lines:
        date_match = re.search(r'\b(\d{2}\.\d{2}\.\d{4})\b', line)
        if date_match:
            date = date_match.group(1)
            if not date.startswith('2028') and 'годности' not in line.lower():
                result['дата_приемки'] = date
                break

    # 3) Название поставщика
    for line in lines:
        supplier_match = re.search(r'(ООО|ЗАО|ОАО|АО|ИП)\s*["«]?([^"«»]+)["»]?', line)
        if supplier_match:
            if 'ВЕРОФАРМ' not in line.upper() and 'ЗАВОД' not in line.upper():
                full_match = re.search(
                    r'(ООО\s*["«]?[^"»]+["»]?|ЗАО\s*["«]?[^"»]+["»]?|ОАО\s*["«]?[^"»]+["»]?|АО\s*["«]?[^"»]+["»]?|ИП\s*["«]?[^"»]+["»]?)',
                    line)
                if full_match:
                    result['название_поставщика'] = full_match.group(1)
                    break

    # 4) ПАРСИНГ ТОВАРОВ
    # Находим начало таблицы
    start_index = -1
    for i, line in enumerate(lines):
        if 'Материальные ценности' in line:
            start_index = i
            break

    # Находим конец таблицы
    end_index = -1
    for i, line in enumerate(lines):
        if line == 'Итого' or 'Итого' in line:
            if i + 1 < len(lines) and (lines[i + 1] == 'X' or lines[i + 1] == 'x'):
                end_index = i + 1
                break
            elif 'X' in line or 'x' in line:
                end_index = i
                break

    if end_index == -1:
        for i, line in enumerate(lines):
            if line == 'X' or line == 'x' or line.startswith('X'):
                if i > 0 and ('Итого' in lines[i - 1] or 'Итого' in lines[i]):
                    end_index = i
                    break

    if start_index == -1 or end_index == -1:
        return result

    # Собираем строки таблицы
    table_lines = lines[start_index:end_index + 1]

    # Находим маркер "14"
    marker_index = -1
    for i, line in enumerate(table_lines):
        if line == '14' or re.search(r'\b14\b', line):
            marker_index = i
            break

    if marker_index == -1:
        for i, line in enumerate(table_lines):
            if re.search(r'\b1\s+2\s+3\s+4\s+5\s+6\s+7\s+8\s+9\s+10\s+11\s+12\s+13\s+14\b', line):
                marker_index = i
                break

    if marker_index == -1:
        return result

    # ПАРСИНГ ТОВАРОВ С ЖЕСТКИМИ ПРАВИЛАМИ
    products = []
    i = 0

    while i < len(table_lines):
        line = table_lines[i]

        # Пропускаем служебные строки
        if any(keyword in line.lower() for keyword in
               ['итого', 'x', 'принял', 'сдал', 'должность', 'подпись', 'материальные ценности', 'измерения',
                'количество', 'цена']):
            i += 1
            continue

        # Ищем номенклатурный номер (SAP) - 8 цифр
        nomencl_match = re.search(r'\b(\d{8})\b', line)
        if nomencl_match:
            nomencl_number = nomencl_match.group(1)

            # Проверяем, что это не служебный номер
            if (not nomencl_number.startswith(('8', '7')) and
                    nomencl_number not in ['48797491', '21000101', '0315003']):

                sap_row = i

                # --- 1. НАЗВАНИЕ ---
                name_parts = []
                for j in range(sap_row - 1, marker_index, -1):
                    prev_line = table_lines[j]
                    if (re.search(r'[А-ЯЁA-Z]', prev_line) and
                            not re.match(r'^\d+$', prev_line) and
                            not any(keyword in prev_line.lower() for keyword in ['итого', 'x', 'принял', 'сдал'])):
                        name_parts.insert(0, prev_line)
                    else:
                        break
                product_name = ' '.join(name_parts).strip() if name_parts else 'Не указано'

                # --- 2. КОЛИЧЕСТВО (SAP + 4 строки) ---
                quantity = None
                qty_row = sap_row + 4
                if qty_row < len(table_lines):
                    qty_line = table_lines[qty_row]
                    qty_match = re.search(r'(\d{1,3}(?:\s?\d{3})*)', qty_line)
                    if qty_match:
                        qty_str = qty_match.group(1).replace(' ', '')
                        try:
                            quantity = int(qty_str)
                        except:
                            pass

                # --- 3. АНАЛИТИЧЕСКИЙ ЛИСТ (SAP + 7 строк) ---
                analytic_list = None
                analytic_row = sap_row + 7
                if analytic_row < len(table_lines):
                    analytic_line = table_lines[analytic_row]
                    analytic_match = re.search(r'\b(1\d{9})\b', analytic_line)
                    if analytic_match:
                        analytic_list = analytic_match.group(1)
                    else:
                        alt_match = re.search(r'\b(\d{10})\b', analytic_line)
                        if alt_match:
                            potential = alt_match.group(1)
                            if not potential.startswith('49'):
                                analytic_list = potential

                # --- 4. ПАРТИЯ ПОСТАВЩИКА (SAP + 8 строк) ---
                supplier_batch = None
                batch_row = sap_row + 8
                if batch_row < len(table_lines):
                    batch_line = table_lines[batch_row]
                    clean_batch = re.sub(r'\d{2}\.\d{2}\.\d{4}', '', batch_line).strip()
                    if clean_batch:
                        supplier_batch = clean_batch
                    else:
                        if not re.match(r'^\d{2}\.\d{2}\.\d{4}$', batch_line):
                            supplier_batch = batch_line

                # --- 5. СРОК ГОДНОСТИ (SAP + 9 строк) ---
                expiration_date = None
                exp_row = sap_row + 9
                if exp_row < len(table_lines):
                    exp_line = table_lines[exp_row]
                    exp_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', exp_line)
                    if exp_match:
                        exp_date = exp_match.group(1)
                        if exp_date != result.get('дата_приемки'):
                            expiration_date = exp_date

                # Добавляем товар
                product_data = {
                    'наименование': product_name,
                    'номенклатурный_номер': nomencl_number,
                    'аналитический_лист': analytic_list if analytic_list else 'Не найден',
                    'партия_поставщика': supplier_batch if supplier_batch else 'Не найдена',
                    'срок_годности': expiration_date if expiration_date else 'Не найден',
                    'количество': quantity if quantity else 'Не найдено'
                }
                products.append(product_data)

        i += 1

    result['товары'] = products
    return result


# ============================================================
# БАЗА ДАННЫХ МАСТЕР-ДАННЫХ МАТЕРИАЛОВ
# ============================================================

class MaterialsDatabase:
    """Класс для работы с базой мастер-данных материалов"""

    def __init__(self, db_file=MATERIALS_DATABASE_FILE):
        self.db_file = db_file
        self.init_database()

    @contextmanager
    def get_connection(self):
        """Получает соединение с базой данных"""
        conn = sqlite3.connect(self.db_file)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
        finally:
            conn.close()

    def init_database(self):
        """Инициализирует базу данных и создает таблицу материалов"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS materials (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    sap_code INTEGER UNIQUE NOT NULL,
                    material_name TEXT UNIQUE NOT NULL,
                    palletization INTEGER DEFAULT 0,
                    box_quantity INTEGER DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            conn.commit()

            cursor.execute("SELECT COUNT(*) FROM materials")
            count = cursor.fetchone()[0]

            if count == 0:
                self.create_default_materials()

    def create_default_materials(self):
        """Создает тестовые материалы"""
        default_materials = [
            {'sap_code': 10000001, 'material_name': 'Тестовый материал 1', 'palletization': 100, 'box_quantity': 10},
            {'sap_code': 10000002, 'material_name': 'Тестовый материал 2', 'palletization': 200, 'box_quantity': 20},
            {'sap_code': 10000003, 'material_name': 'Тестовый материал 3', 'palletization': 150, 'box_quantity': 15},
        ]

        for material in default_materials:
            self.add_material(material)

    def add_material(self, material_data):
        """Добавляет новый материал в базу"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO materials (sap_code, material_name, palletization, box_quantity)
                    VALUES (?, ?, ?, ?)
                ''', (
                    material_data['sap_code'],
                    material_data['material_name'],
                    material_data.get('palletization', 0),
                    material_data.get('box_quantity', 0)
                ))
                conn.commit()
                return True, "Материал успешно добавлен"
        except sqlite3.IntegrityError as e:
            if 'UNIQUE constraint failed: materials.sap_code' in str(e):
                return False, f"Материал с SAP-кодом {material_data['sap_code']} уже существует"
            elif 'UNIQUE constraint failed: materials.material_name' in str(e):
                return False, f"Материал с названием '{material_data['material_name']}' уже существует"
            return False, f"Ошибка целостности данных: {str(e)}"
        except Exception as e:
            return False, f"Ошибка при добавлении материала: {str(e)}"

    def get_material_by_sap(self, sap_code):
        """Находит материал по SAP-коду"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM materials WHERE sap_code = ?", (sap_code,))
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None

    def get_material_by_name(self, material_name):
        """Находит материал по названию"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM materials WHERE material_name = ?", (material_name,))
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None

    def update_material(self, sap_code, new_data):
        """Обновляет данные материала"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()

                fields = []
                values = []

                for key in ['material_name', 'palletization', 'box_quantity']:
                    if key in new_data:
                        fields.append(f"{key} = ?")
                        values.append(new_data[key])

                values.append(sap_code)
                fields.append("updated_at = CURRENT_TIMESTAMP")

                query = f"UPDATE materials SET {', '.join(fields)} WHERE sap_code = ?"
                cursor.execute(query, values)
                conn.commit()
                return True, "Данные материала обновлены"
        except sqlite3.IntegrityError as e:
            if 'UNIQUE constraint failed: materials.material_name' in str(e):
                return False, f"Материал с названием '{new_data.get('material_name', '')}' уже существует"
            return False, f"Ошибка целостности данных: {str(e)}"
        except Exception as e:
            return False, f"Ошибка при обновлении: {str(e)}"

    def delete_material(self, sap_code):
        """Удаляет материал"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("DELETE FROM materials WHERE sap_code = ?", (sap_code,))
                conn.commit()
                return True, "Материал удален"
        except Exception as e:
            return False, f"Ошибка при удалении: {str(e)}"

    def get_all_materials(self):
        """Возвращает список всех материалов"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM materials ORDER BY sap_code")
            rows = cursor.fetchall()
            return [dict(row) for row in rows]

    def get_materials_count(self):
        """Возвращает количество материалов в базе"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM materials")
            return cursor.fetchone()[0]

    def validate_and_import_from_excel(self, filepath):
        """Проверяет и импортирует данные из Excel файла (без pandas)"""
        try:
            wb = load_workbook(filepath, data_only=True)
            ws = wb.active

            is_header = False
            first_row = [cell.value for cell in ws[1]]
            if first_row and any(isinstance(cell, str) and not str(cell).isdigit() for cell in first_row[:2]):
                is_header = True
                start_row = 2
            else:
                start_row = 1

            errors = []
            materials_to_add = []
            materials_to_update = []

            for row_idx in range(start_row, ws.max_row + 1):
                row = [ws.cell(row=row_idx, column=col).value for col in range(1, 4)]

                sap_value = row[0]
                if sap_value is None:
                    errors.append(f"Строка {row_idx}: SAP код не может быть пустым")
                    continue

                try:
                    sap_code = int(sap_value)
                    if len(str(sap_code)) != 8:
                        errors.append(
                            f"Строка {row_idx}: SAP код должен содержать 8 цифр (текущее значение: {sap_code})")
                        continue
                except (ValueError, TypeError):
                    errors.append(f"Строка {row_idx}: SAP код должен быть числом (текущее значение: {sap_value})")
                    continue

                material_name = str(row[1]) if row[1] else ''
                if not material_name.strip():
                    errors.append(f"Строка {row_idx}: Название материала не может быть пустым")
                    continue

                palletization_value = row[2] if row[2] else 0
                try:
                    palletization = int(palletization_value)
                    if palletization < 0:
                        errors.append(
                            f"Строка {row_idx}: Палетизация не может быть отрицательной (текущее значение: {palletization})")
                        continue
                except (ValueError, TypeError):
                    errors.append(
                        f"Строка {row_idx}: Палетизация должна быть числом (текущее значение: {palletization_value})")
                    continue

                if ws.max_column > 3:
                    for col_idx in range(4, ws.max_column + 1):
                        if ws.cell(row=row_idx, column=col_idx).value is not None:
                            errors.append(f"Строка {row_idx}: Обнаружены данные в дополнительной колонке {col_idx}")
                            break

                material_data = {
                    'sap_code': sap_code,
                    'material_name': material_name.strip(),
                    'palletization': palletization
                }

                existing = self.get_material_by_sap(sap_code)
                if existing:
                    materials_to_update.append(material_data)
                else:
                    existing_name = self.get_material_by_name(material_name.strip())
                    if existing_name:
                        errors.append(
                            f"Строка {row_idx}: Материал с названием '{material_name.strip()}' уже существует в базе (SAP: {existing_name['sap_code']})")
                    else:
                        materials_to_add.append(material_data)

            if errors:
                return False, "Ошибки в файле:\n" + "\n".join(errors)

            updated_count = 0
            for material in materials_to_update:
                existing = self.get_material_by_sap(material['sap_code'])
                if existing:
                    new_data = {
                        'material_name': material['material_name'],
                        'palletization': material['palletization'],
                        'box_quantity': existing.get('box_quantity', 0)
                    }
                    success, message = self.update_material(material['sap_code'], new_data)
                    if success:
                        updated_count += 1

            added_count = 0
            for material in materials_to_add:
                material_data = {
                    'sap_code': material['sap_code'],
                    'material_name': material['material_name'],
                    'palletization': material['palletization'],
                    'box_quantity': 0
                }
                success, message = self.add_material(material_data)
                if success:
                    added_count += 1

            return True, f"Успешно обработано: добавлено {added_count} материалов, обновлено {updated_count} материалов"

        except Exception as e:
            print(f"❌ Ошибка при импорте: {e}")
            traceback.print_exc()
            return False, f"Ошибка при обработке файла: {str(e)}"

    def export_to_excel(self):
        """Экспортирует данные материалов в Excel файл"""
        try:
            materials = self.get_all_materials()

            wb = Workbook()
            ws = wb.active
            ws.title = "Реестр сырья и материалов"

            # Заголовки
            headers = ['Номер SAP', 'Наименование', 'Палетизация', 'Вложение в короб']
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="1a237e", end_color="1a237e", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")

            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment

            # Данные
            for row_idx, material in enumerate(materials, 2):
                ws.cell(row=row_idx, column=1, value=material['sap_code'])
                ws.cell(row=row_idx, column=2, value=material['material_name'])
                ws.cell(row=row_idx, column=3, value=material['palletization'])
                ws.cell(row=row_idx, column=4, value=material['box_quantity'])

            # Настройка ширины колонок
            ws.column_dimensions['A'].width = 15
            ws.column_dimensions['B'].width = 40
            ws.column_dimensions['C'].width = 15
            ws.column_dimensions['D'].width = 20

            # Сохраняем файл
            filename = f"Реестр_сырья_и_материалов_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            wb.save(filepath)

            return filepath

        except Exception as e:
            print(f"❌ Ошибка экспорта: {e}")
            traceback.print_exc()
            return None


# ============================================================
# ШИФРОВАННАЯ БАЗА ДАННЫХ ПОЛЬЗОВАТЕЛЕЙ (SQLite)
# ============================================================

class UserDatabase:
    """Класс для работы с зашифрованной базой данных пользователей"""

    def __init__(self, db_file=DATABASE_FILE):
        self.db_file = db_file
        self.master_key = self._get_master_key()
        self.init_database()

    def _get_master_key(self):
        """Получает или создает мастер-ключ для шифрования"""
        key_file = os.path.join(BASE_DIR, 'master.key')
        if os.path.exists(key_file):
            with open(key_file, 'rb') as f:
                return f.read()
        else:
            key = secrets.token_bytes(32)
            with open(key_file, 'wb') as f:
                f.write(key)
            return key

    def _encrypt(self, data):
        """Шифрует данные"""
        if not data:
            return ''
        key = self.master_key
        encrypted = bytearray()
        for i, char in enumerate(str(data).encode('utf-8')):
            encrypted.append(char ^ key[i % len(key)])
        return encrypted.hex()

    def _decrypt(self, encrypted_data):
        """Расшифровывает данные"""
        if not encrypted_data:
            return ''
        key = self.master_key
        encrypted_bytes = bytes.fromhex(encrypted_data)
        decrypted = bytearray()
        for i, byte in enumerate(encrypted_bytes):
            decrypted.append(byte ^ key[i % len(key)])
        return decrypted.decode('utf-8')

    @contextmanager
    def get_connection(self):
        """Получает соединение с базой данных"""
        conn = sqlite3.connect(self.db_file)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
        finally:
            conn.close()

    def init_database(self):
        """Инициализирует базу данных и создает таблицы"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    login TEXT UNIQUE NOT NULL,
                    email TEXT NOT NULL,
                    password_hash TEXT NOT NULL,
                    fname TEXT,
                    lname TEXT,
                    enterprise TEXT,
                    plant TEXT,
                    rights TEXT,
                    status TEXT DEFAULT 'active',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            cursor.execute("PRAGMA table_info(users)")
            columns = [col[1] for col in cursor.fetchall()]
            if 'status' not in columns:
                cursor.execute("ALTER TABLE users ADD COLUMN status TEXT DEFAULT 'active'")
                conn.commit()
                cursor.execute("UPDATE users SET status = 'active' WHERE status IS NULL")
                conn.commit()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS audit_log (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_login TEXT,
                    action TEXT,
                    details TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            conn.commit()

            cursor.execute("SELECT COUNT(*) FROM users")
            count = cursor.fetchone()[0]

            if count == 0:
                self.create_default_users()

    def create_default_users(self):
        """Создает пользователей по умолчанию"""
        default_users = [
            {
                'login': 'admin',
                'email': 'admin@example.com',
                'password': 'admin123',
                'fname': 'Администратор',
                'lname': 'Системы',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Администратор',
                'status': 'active'
            },
            {
                'login': 'operator',
                'email': 'operator@example.com',
                'password': 'operator123',
                'fname': 'Оператор',
                'lname': 'СиМ',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Оператор СиМ',
                'status': 'active'
            },
            {
                'login': 'alexander.abramov@abbott.com',
                'email': 'alexander.abramov@abbott.com',
                'password': 'alexander.abramov@abbott.com',
                'fname': 'Александр',
                'lname': 'Абрамов',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Администратор',
                'status': 'active'
            },
            {
                'login': 'anastasia.yudashkina@abbott.com',
                'email': 'anastasia.yudashkina@abbott.com',
                'password': 'anastasia.yudashkina@abbott.com',
                'fname': 'Анастасия',
                'lname': 'Юдашкина',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Оператор СиМ',
                'status': 'active'
            }
        ]

        for user in default_users:
            self.add_user(user)

    def add_user(self, user_data):
        """Добавляет нового пользователя в базу"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                password_hash = generate_password_hash(user_data['password'])

                cursor.execute('''
                    INSERT INTO users (login, email, password_hash, fname, lname, enterprise, plant, rights, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    user_data['login'],
                    user_data['email'],
                    password_hash,
                    user_data.get('fname', ''),
                    user_data.get('lname', ''),
                    user_data.get('enterprise', ''),
                    user_data.get('plant', ''),
                    user_data.get('rights', 'Оператор СиМ'),
                    user_data.get('status', 'active')
                ))

                conn.commit()
                self.log_action(user_data['login'], 'ADD_USER', f"Добавлен пользователь {user_data['login']}")
                return True, "Пользователь успешно добавлен"
        except sqlite3.IntegrityError:
            return False, "Пользователь с таким логином уже существует"
        except Exception as e:
            return False, f"Ошибка при добавлении пользователя: {str(e)}"

    def find_user(self, login):
        """Находит пользователя по логину"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM users WHERE login = ?", (login,))
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None

    def authenticate(self, login, password):
        """Проверяет логин и пароль"""
        user = self.find_user(login)
        if user:
            if user.get('status') == 'blocked':
                self.log_action(login, 'LOGIN_FAILED', f"Попытка входа заблокированного пользователя")
                return None, 'blocked'
            if user.get('status') == 'deleted':
                self.log_action(login, 'LOGIN_FAILED', f"Попытка входа удаленного пользователя")
                return None, 'deleted'

            if check_password_hash(user['password_hash'], password):
                self.log_action(login, 'LOGIN_SUCCESS', f"Успешный вход")
                return user, 'success'
            else:
                self.log_action(login, 'LOGIN_FAILED', f"Неверный пароль")
                return None, 'invalid'

    def update_password(self, login, new_password):
        """Обновляет пароль пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                password_hash = generate_password_hash(new_password)
                cursor.execute(
                    "UPDATE users SET password_hash = ?, updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (password_hash, login)
                )
                conn.commit()
                self.log_action(login, 'UPDATE_PASSWORD', "Пароль изменен")
                return True
        except Exception as e:
            print(f"Ошибка обновления пароля: {e}")
            return False

    def update_user(self, login, new_data):
        """Обновляет данные пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()

                fields = []
                values = []

                for key in ['email', 'fname', 'lname', 'enterprise', 'plant', 'rights', 'status']:
                    if key in new_data:
                        fields.append(f"{key} = ?")
                        values.append(new_data[key])

                if 'password' in new_data and new_data['password']:
                    fields.append("password_hash = ?")
                    values.append(generate_password_hash(new_data['password']))

                values.append(login)
                fields.append("updated_at = CURRENT_TIMESTAMP")

                query = f"UPDATE users SET {', '.join(fields)} WHERE login = ?"
                cursor.execute(query, values)
                conn.commit()

                self.log_action(login, 'UPDATE_USER', f"Данные пользователя обновлены")
                return True, "Данные пользователя обновлены"
        except Exception as e:
            return False, f"Ошибка обновления: {str(e)}"

    def delete_user(self, login):
        """Удаляет пользователя (soft delete - меняет статус)"""
        admins = self.get_users_by_rights('Администратор')
        user = self.find_user(login)

        if user and user['rights'] == 'Администратор' and len(admins) <= 1:
            return False, "Нельзя удалить единственного администратора"

        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE users SET status = 'deleted', updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (login,)
                )
                conn.commit()
                self.log_action(login, 'DELETE_USER', f"Пользователь удален")
                return True, "Пользователь удален"
        except Exception as e:
            return False, f"Ошибка удаления: {str(e)}"

    def block_user(self, login):
        """Блокирует пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE users SET status = 'blocked', updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (login,)
                )
                conn.commit()
                self.log_action(login, 'BLOCK_USER', f"Пользователь заблокирован")
                return True, "Пользователь заблокирован"
        except Exception as e:
            return False, f"Ошибка блокировки: {str(e)}"

    def unblock_user(self, login):
        """Разблокирует пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE users SET status = 'active', updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (login,)
                )
                conn.commit()
                self.log_action(login, 'UNBLOCK_USER', f"Пользователь разблокирован")
                return True, "Пользователь разблокирован"
        except Exception as e:
            return False, f"Ошибка разблокировки: {str(e)}"

    def get_all_users(self):
        """Возвращает список всех пользователей"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id, login, email, fname, lname, enterprise, plant, rights, status, created_at FROM users ORDER BY login"
            )
            rows = cursor.fetchall()
            return [dict(row) for row in rows]

    def get_users_by_rights(self, rights):
        """Возвращает пользователей с определенными правами"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM users WHERE rights = ? AND status = 'active'", (rights,))
            rows = cursor.fetchall()
            return [dict(row) for row in rows]

    def log_action(self, user_login, action, details):
        """Логирует действие пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO audit_log (user_login, action, details)
                    VALUES (?, ?, ?)
                ''', (user_login, action, details))
                conn.commit()
        except Exception as e:
            print(f"Ошибка логирования: {e}")

    def get_audit_log(self, limit=100):
        """Возвращает последние записи аудита"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM audit_log ORDER BY timestamp DESC LIMIT ?",
                (limit,)
            )
            rows = cursor.fetchall()
            return [dict(row) for row in rows]


# ============================================================
# ИНИЦИАЛИЗАЦИЯ БАЗ ДАННЫХ
# ============================================================

db = UserDatabase(DATABASE_FILE)
materials_db = MaterialsDatabase(MATERIALS_DATABASE_FILE)


# ============================================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С НАСТРОЙКАМИ
# ============================================================

def load_settings():
    """Загружает настройки из файла"""
    default_settings = {
        'save_path': os.path.join(BASE_DIR, 'generated_labels'),
        'sop_code': 'RUPD.VLG.05.04.C01',
        'font_settings': {
            'header_size': 12,
            'product_size_short': 80,
            'product_size_medium': 72,
            'product_size_long': 64,
            'supplier_size_short': 42,
            'supplier_size_medium': 42,
            'supplier_size_long': 36,
            'analytic_size': 117,
            'code_size': 14,
            'line_spacing_short': 1.5,
            'line_spacing_medium': 1.0,
            'line_spacing_long': 0.8,
            'short_name_threshold': 25,
            'medium_name_threshold': 40
        }
    }

    try:
        if os.path.exists(SETTINGS_FILE):
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                settings = json.load(f)
            if 'save_path' not in settings:
                settings['save_path'] = default_settings['save_path']
            if 'sop_code' not in settings:
                settings['sop_code'] = default_settings['sop_code']
            if 'font_settings' not in settings:
                settings['font_settings'] = default_settings['font_settings']
            else:
                for key in default_settings['font_settings']:
                    if key not in settings['font_settings']:
                        settings['font_settings'][key] = default_settings['font_settings'][key]
            return settings
        else:
            os.makedirs(default_settings['save_path'], exist_ok=True)
            save_settings(default_settings)
            return default_settings
    except Exception as e:
        print(f"Ошибка загрузки настроек: {e}")
        return default_settings


def save_settings(settings):
    """Сохраняет настройки в файл"""
    try:
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
        return True
    except Exception as e:
        print(f"Ошибка сохранения настроек: {e}")
        return False


def get_available_drives():
    """Получает список доступных дисков в Windows"""
    drives = []
    if sys.platform == 'win32':
        import string
        for letter in string.ascii_uppercase:
            drive = f"{letter}:\\"
            if os.path.exists(drive):
                drives.append(drive)
    else:
        drives = ['/']
    return drives


def get_subdirectories(path):
    """Получает список подпапок в указанном пути"""
    subdirs = []
    try:
        if os.path.exists(path):
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                if os.path.isdir(item_path):
                    subdirs.append(item)
    except:
        pass
    return sorted(subdirs)


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================

def calculate_pallets(quantity, palletization):
    """Рассчитывает количество паллет"""
    if not quantity:
        return 1

    try:
        qty = int(quantity)
    except:
        return 1

    if palletization > 0:
        try:
            return math.ceil(qty / palletization)
        except:
            return 1
    else:
        return 1


def clean_filename(text):
    """Очищает текст от недопустимых символов для имени файла"""
    if not text:
        return 'Неизвестный_поставщик'
    invalid_chars = r'[<>:"/\\|?*]'
    cleaned = re.sub(invalid_chars, '', text)
    cleaned = ' '.join(cleaned.split())
    if not cleaned.strip():
        return 'Неизвестный_поставщик'
    return cleaned


def get_unique_filename(base_path, base_name, extension='.docx'):
    """Формирует уникальное имя файла с индексацией при повторении"""
    counter = 1
    while True:
        if counter == 1:
            filename = f"{base_name}{extension}"
        else:
            filename = f"{base_name}_{counter}{extension}"

        full_path = os.path.join(base_path, filename)
        if not os.path.exists(full_path):
            return full_path
        counter += 1


# ============================================================
# ФУНКЦИЯ ГЕНЕРАЦИИ ЭТИКЕТОК
# ============================================================

def get_name_length_category(name_len, font_settings):
    """Определяет категорию длины названия"""
    short_threshold = font_settings.get('short_name_threshold', 25)
    medium_threshold = font_settings.get('medium_name_threshold', 40)

    if name_len < short_threshold:
        return 'short'
    elif name_len < medium_threshold:
        return 'medium'
    else:
        return 'long'


def get_font_and_spacing_settings(product_name, font_settings):
    """Определяет размеры шрифтов и межстрочный интервал"""
    name_len = len(product_name)
    category = get_name_length_category(name_len, font_settings)

    if category == 'short':
        return {
            'header': font_settings.get('header_size', 12),
            'product': font_settings.get('product_size_short', 80),
            'supplier': font_settings.get('supplier_size_short', 42),
            'analytic': font_settings.get('analytic_size', 117),
            'code': font_settings.get('code_size', 14),
            'line_spacing': font_settings.get('line_spacing_short', 1.5),
            'category': 'short'
        }
    elif category == 'medium':
        return {
            'header': font_settings.get('header_size', 12),
            'product': font_settings.get('product_size_medium', 72),
            'supplier': font_settings.get('supplier_size_medium', 42),
            'analytic': font_settings.get('analytic_size', 117),
            'code': font_settings.get('code_size', 14),
            'line_spacing': font_settings.get('line_spacing_medium', 1.0),
            'category': 'medium'
        }
    else:
        return {
            'header': font_settings.get('header_size', 12),
            'product': font_settings.get('product_size_long', 64),
            'supplier': font_settings.get('supplier_size_long', 36),
            'analytic': font_settings.get('analytic_size', 117),
            'code': font_settings.get('code_size', 14),
            'line_spacing': font_settings.get('line_spacing_long', 0.8),
            'category': 'long'
        }


def create_pallet_labels_file(all_products, order_number, supplier_name, order_date, selected_indices, save_path=None,
                              sop_code=None):
    """Создает файл с палетными этикетками только для выбранных товаров"""
    if save_path is None:
        settings = load_settings()
        save_path = settings.get('save_path', os.path.join(BASE_DIR, 'generated_labels'))

    if sop_code is None:
        settings = load_settings()
        sop_code = settings.get('sop_code', 'RUPD.VLG.05.04.C01')

    settings = load_settings()
    font_settings = settings.get('font_settings', {})

    try:
        os.makedirs(save_path, exist_ok=True)
    except Exception as e:
        save_path = os.path.join(BASE_DIR, 'generated_labels')
        os.makedirs(save_path, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.15)
        section.bottom_margin = Inches(0.15)
        section.left_margin = Inches(0.15)
        section.right_margin = Inches(0.15)

    total_labels = 0
    first_page = True

    for idx, product in enumerate(all_products):
        if idx not in selected_indices:
            continue

        pallets_count = product.get('pallets', 1)
        if pallets_count < 1:
            pallets_count = 1

        product_name = product.get('наименование', 'Не указано')
        supplier_series = product.get('партия_поставщика', 'Не указана')
        analytic_list = product.get('аналитический_лист', 'Не указан')

        font_settings_dict = get_font_and_spacing_settings(product_name, font_settings)
        line_spacing = font_settings_dict.get('line_spacing', 1.0)

        for pallet_num in range(pallets_count):
            if not first_page:
                doc.add_page_break()
                if doc.sections:
                    section = doc.sections[-1]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11.69)
                    section.page_height = Inches(8.27)
                    section.top_margin = Inches(0.15)
                    section.bottom_margin = Inches(0.15)
                    section.left_margin = Inches(0.15)
                    section.right_margin = Inches(0.15)
            else:
                first_page = False

            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(10.8)

            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].clear()

            p1 = cell.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(2)
            p1.paragraph_format.space_after = Pt(2)
            p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p1.paragraph_format.line_spacing = line_spacing
            run1 = p1.add_run('ООО «Верофарм»')
            run1.italic = True
            run1.bold = True
            run1.font.size = Pt(font_settings_dict['header'])
            run1.font.name = 'Calibri'

            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(5)
            p2.paragraph_format.space_after = Pt(5)
            p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p2.paragraph_format.line_spacing = line_spacing
            run2 = p2.add_run(product_name)
            run2.italic = True
            run2.bold = True
            run2.font.size = Pt(font_settings_dict['product'])
            run2.font.name = 'Calibri'

            supplier_text = f"п. {supplier_series}" if supplier_series and supplier_series != 'Не указана' else supplier_series
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(4)
            p3.paragraph_format.space_after = Pt(4)
            p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p3.paragraph_format.line_spacing = line_spacing
            run3 = p3.add_run(supplier_text)
            run3.italic = True
            run3.bold = False
            run3.font.size = Pt(font_settings_dict['supplier'])
            run3.font.name = 'Calibri'

            p4 = cell.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(4)
            p4.paragraph_format.space_after = Pt(4)
            p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p4.paragraph_format.line_spacing = line_spacing
            run4 = p4.add_run(f'АЛ {analytic_list}')
            run4.italic = True
            run4.bold = True
            run4.font.size = Pt(font_settings_dict['analytic'])
            run4.font.name = 'Calibri'

            p5 = cell.add_paragraph()
            p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p5.paragraph_format.space_before = Pt(8)
            p5.paragraph_format.space_after = Pt(2)
            p5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p5.paragraph_format.line_spacing = line_spacing
            run5 = p5.add_run(sop_code)
            run5.italic = False
            run5.bold = True
            run5.font.size = Pt(font_settings_dict['code'])
            run5.font.name = 'Calibri'

            total_labels += 1

    clean_supplier = clean_filename(supplier_name) if supplier_name else 'Неизвестный_поставщик'

    date_str = ''
    if order_date:
        date_match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', order_date)
        if date_match:
            date_str = f"{date_match.group(1)}{date_match.group(2)}{date_match.group(3)}"
        else:
            date_str = datetime.now().strftime("%d%m%Y")
    else:
        date_str = datetime.now().strftime("%d%m%Y")

    base_name = f"{clean_supplier}_{order_number}_{date_str}"
    filepath = get_unique_filename(save_path, base_name, '.docx')

    try:
        doc.save(filepath)
    except Exception as e:
        temp_dir = tempfile.gettempdir()
        filepath = get_unique_filename(temp_dir, base_name, '.docx')
        doc.save(filepath)

    return filepath, total_labels


def open_in_word(filepath):
    """Открывает файл в Microsoft Word"""
    try:
        if sys.platform == 'win32':
            os.startfile(filepath)
        elif sys.platform == 'darwin':
            subprocess.run(['open', filepath])
        else:
            subprocess.run(['xdg-open', filepath])
        return True
    except Exception as e:
        return False


# ============================================================
# ДЕКОРАТОРЫ ДЛЯ ПРОВЕРКИ ПРАВ ДОСТУПА
# ============================================================

def login_required(f):
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


def admin_required(f):
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))
        if session.get('user_rights') != 'Администратор':
            flash('Доступ запрещен. Требуются права администратора.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


def manager_required(f):
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))
        user_rights = session.get('user_rights')
        if user_rights not in ['Администратор', 'Менеджер ОСЛ']:
            flash('Доступ запрещен. Требуются права Администратора или Менеджера ОСЛ.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


def labels_access_required(f):
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))

        user_rights = session.get('user_rights')
        user_plant = session.get('user_plant')

        has_access = False

        if user_rights in ['Администратор', 'Оператор СиМ', 'Менеджер ОСЛ'] and user_plant == 'Вольгинский':
            has_access = True

        if not has_access:
            flash('Доступ запрещен. У вас недостаточно прав для формирования этикеток.', 'danger')
            return redirect(url_for('index'))

        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


# ============================================================
# МАРШРУТЫ ПРИЛОЖЕНИЯ
# ============================================================

@app.route('/')
def index():
    try:
        if 'user' in session:
            return render_template('index.html')
        return redirect(url_for('login'))
    except Exception as e:
        print(f"Ошибка в index: {e}")
        traceback.print_exc()
        return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    try:
        if request.method == 'POST':
            login = request.form.get('login', '').strip()
            password = request.form.get('password', '')
            remember = request.form.get('remember', False)

            if not login or not password:
                flash('Введите логин и пароль', 'warning')
                return render_template('login.html')

            user, status = db.authenticate(login, password)

            if status == 'blocked':
                flash('Пользователь заблокирован. Обратитесь к администратору.', 'danger')
                return render_template('login.html')

            if status == 'deleted':
                flash('Пользователь удален. Обратитесь к администратору.', 'danger')
                return render_template('login.html')

            if user:
                session['user'] = user['login']
                session['user_email'] = user['email']
                session['user_fname'] = user['fname']
                session['user_lname'] = user['lname']
                session['user_enterprise'] = user['enterprise']
                session['user_plant'] = user['plant']
                session['user_rights'] = user['rights']

                if remember:
                    session.permanent = True

                flash(f'Добро пожаловать, {user["fname"]} {user["lname"]}!', 'success')
                return redirect(url_for('index'))
            else:
                flash('Неверный логин или пароль. Проверьте правильность ввода.', 'danger')

        return render_template('login.html')
    except Exception as e:
        print(f"Ошибка в login: {e}")
        traceback.print_exc()
        flash('Ошибка при входе в систему', 'danger')
        return render_template('login.html')


@app.route('/logout')
def logout():
    if 'user' in session:
        db.log_action(session['user'], 'LOGOUT', "Выход из системы")
    session.clear()
    flash('Вы вышли из системы', 'info')
    return redirect(url_for('login'))


@app.route('/change_password', methods=['GET', 'POST'])
@login_required
def change_password():
    try:
        if request.method == 'POST':
            current_password = request.form.get('current_password', '')
            new_password = request.form.get('new_password', '')
            confirm_password = request.form.get('confirm_password', '')

            user = db.find_user(session['user'])

            if not user:
                flash('Пользователь не найден', 'danger')
                return redirect(url_for('index'))

            if not check_password_hash(user['password_hash'], current_password):
                flash('Неверный текущий пароль', 'danger')
                return render_template('change_password.html')

            if len(new_password) < 6:
                flash('Новый пароль должен содержать минимум 6 символов', 'warning')
                return render_template('change_password.html')

            if new_password != confirm_password:
                flash('Пароли не совпадают', 'warning')
                return render_template('change_password.html')

            if db.update_password(session['user'], new_password):
                flash('Пароль успешно изменен', 'success')
                return redirect(url_for('index'))
            else:
                flash('Ошибка при смене пароля', 'danger')

        return render_template('change_password.html')
    except Exception as e:
        print(f"Ошибка в change_password: {e}")
        traceback.print_exc()
        flash('Ошибка при смене пароля', 'danger')
        return redirect(url_for('index'))


@app.route('/admin')
@manager_required
def admin_panel():
    try:
        users = db.get_all_users()
        audit_log = db.get_audit_log(50)
        user_rights = session.get('user_rights')
        return render_template('admin.html', users=users, audit_log=audit_log, user_rights=user_rights)
    except Exception as e:
        print(f"Ошибка в admin_panel: {e}")
        traceback.print_exc()
        flash(f'Ошибка загрузки панели администратора: {str(e)}', 'danger')
        return redirect(url_for('settings_page'))


@app.route('/admin/add_user', methods=['POST'])
@manager_required
def add_user():
    try:
        login = request.form.get('login', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        fname = request.form.get('fname', '').strip()
        lname = request.form.get('lname', '').strip()
        enterprise = request.form.get('enterprise', 'ООО "Верофарм"')
        plant = request.form.get('plant', 'Вольгинский')
        rights = request.form.get('rights', 'Оператор СиМ')
        status = request.form.get('status', 'active')

        user_rights = session.get('user_rights')

        if user_rights == 'Менеджер ОСЛ' and rights != 'Оператор СиМ':
            flash('Менеджер ОСЛ может создавать только пользователей с правами "Оператор СиМ"', 'danger')
            return redirect(url_for('admin_panel'))

        if not login or not email or not password:
            flash('Логин, email и пароль обязательны для заполнения', 'warning')
            return redirect(url_for('admin_panel'))

        if len(password) < 6:
            flash('Пароль должен содержать минимум 6 символов', 'warning')
            return redirect(url_for('admin_panel'))

        user_data = {
            'login': login,
            'email': email,
            'password': password,
            'fname': fname,
            'lname': lname,
            'enterprise': enterprise,
            'plant': plant,
            'rights': rights,
            'status': status
        }

        success, message = db.add_user(user_data)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в add_user: {e}")
        traceback.print_exc()
        flash('Ошибка при добавлении пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/delete_user/<login>', methods=['POST'])
@admin_required
def delete_user(login):
    try:
        if login == session['user']:
            flash('Нельзя удалить самого себя', 'danger')
            return redirect(url_for('admin_panel'))

        success, message = db.delete_user(login)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в delete_user: {e}")
        traceback.print_exc()
        flash('Ошибка при удалении пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/block_user/<login>', methods=['POST'])
@manager_required
def block_user(login):
    try:
        if login == session['user']:
            flash('Нельзя заблокировать самого себя', 'danger')
            return redirect(url_for('admin_panel'))

        user_rights = session.get('user_rights')
        target_user = db.find_user(login)

        if user_rights == 'Менеджер ОСЛ':
            if not target_user or target_user.get('rights') != 'Оператор СиМ':
                flash('Менеджер ОСЛ может блокировать только пользователей с правами "Оператор СиМ"', 'danger')
                return redirect(url_for('admin_panel'))

        success, message = db.block_user(login)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в block_user: {e}")
        traceback.print_exc()
        flash('Ошибка при блокировке пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/unblock_user/<login>', methods=['POST'])
@manager_required
def unblock_user(login):
    try:
        user_rights = session.get('user_rights')
        target_user = db.find_user(login)

        if user_rights == 'Менеджер ОСЛ':
            if not target_user or target_user.get('rights') != 'Оператор СиМ':
                flash('Менеджер ОСЛ может разблокировать только пользователей с правами "Оператор СиМ"', 'danger')
                return redirect(url_for('admin_panel'))

        success, message = db.unblock_user(login)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в unblock_user: {e}")
        traceback.print_exc()
        flash('Ошибка при разблокировке пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/update_user/<login>', methods=['POST'])
@admin_required
def update_user(login):
    try:
        new_data = {
            'email': request.form.get('email', '').strip(),
            'fname': request.form.get('fname', '').strip(),
            'lname': request.form.get('lname', '').strip(),
            'enterprise': request.form.get('enterprise', 'ООО "Верофарм"'),
            'plant': request.form.get('plant', 'Вольгинский'),
            'rights': request.form.get('rights', 'Оператор СиМ'),
            'status': request.form.get('status', 'active')
        }

        new_password = request.form.get('password', '')
        if new_password:
            if len(new_password) < 6:
                flash('Пароль должен содержать минимум 6 символов', 'warning')
                return redirect(url_for('admin_panel'))
            new_data['password'] = new_password

        success, message = db.update_user(login, new_data)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в update_user: {e}")
        traceback.print_exc()
        flash('Ошибка при обновлении пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/get_user_data/<login>')
@manager_required
def get_user_data(login):
    try:
        user = db.find_user(login)
        if user:
            return jsonify({
                'success': True,
                'email': user.get('email', ''),
                'fname': user.get('fname', ''),
                'lname': user.get('lname', ''),
                'enterprise': user.get('enterprise', 'ООО "Верофарм"'),
                'plant': user.get('plant', 'Вольгинский'),
                'rights': user.get('rights', 'Оператор СиМ'),
                'status': user.get('status', 'active')
            })
        return jsonify({'success': False, 'error': 'Пользователь не найден'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/materials')
@manager_required
def materials_page():
    try:
        materials = materials_db.get_all_materials()
        return render_template('materials.html', materials=materials)
    except Exception as e:
        print(f"Ошибка в materials_page: {e}")
        traceback.print_exc()
        flash('Ошибка загрузки данных материалов', 'danger')
        return redirect(url_for('settings_page'))


@app.route('/materials/update', methods=['POST'])
@manager_required
def update_materials():
    try:
        # Получаем данные из формы
        sap_codes = request.form.getlist('sap_code[]')
        material_names = request.form.getlist('material_name[]')
        palletizations = request.form.getlist('palletization[]')
        box_quantities = request.form.getlist('box_quantity[]')

        if not sap_codes:
            flash('Нет данных для сохранения', 'danger')
            return redirect(url_for('materials_page'))

        errors = []
        success_count = 0

        for i in range(len(sap_codes)):
            sap_code = sap_codes[i]
            material_name = material_names[i].strip() if i < len(material_names) else ''
            try:
                palletization = int(palletizations[i]) if i < len(palletizations) and palletizations[i] else 0
            except:
                palletization = 0
            try:
                box_quantity = int(box_quantities[i]) if i < len(box_quantities) and box_quantities[i] else 0
            except:
                box_quantity = 0

            if not material_name:
                errors.append(f"SAP {sap_code}: Название материала не может быть пустым")
                continue

            # Проверяем уникальность названия
            existing = materials_db.get_material_by_name(material_name)
            if existing and str(existing['sap_code']) != str(sap_code):
                errors.append(f"Материал с названием '{material_name}' уже существует (SAP: {existing['sap_code']})")
                continue

            if palletization < 0:
                errors.append(f"SAP {sap_code}: Палетизация не может быть отрицательной")
                continue

            if box_quantity < 0:
                errors.append(f"SAP {sap_code}: Вложение в короб не может быть отрицательным")
                continue

            new_data = {
                'material_name': material_name,
                'palletization': palletization,
                'box_quantity': box_quantity
            }
            success, message = materials_db.update_material(int(sap_code), new_data)
            if success:
                success_count += 1
            else:
                errors.append(message)

        if errors:
            flash(f"Сохранено {success_count} материалов. Ошибки:\n" + "\n".join(errors), 'warning')
        else:
            flash(f"Успешно сохранено {success_count} материалов", 'success')

        return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в update_materials: {e}")
        traceback.print_exc()
        flash(f'Ошибка при сохранении данных: {str(e)}', 'danger')
        return redirect(url_for('materials_page'))


@app.route('/materials/delete/<int:sap_code>', methods=['POST'])
@manager_required
def delete_material(sap_code):
    try:
        success, message = materials_db.delete_material(sap_code)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в delete_material: {e}")
        traceback.print_exc()
        flash(f'Ошибка при удалении материала: {str(e)}', 'danger')
        return redirect(url_for('materials_page'))


@app.route('/materials/import', methods=['GET', 'POST'])
@manager_required
def import_materials():
    if request.method == 'GET':
        return render_template('import_materials.html')

    try:
        if 'file' not in request.files:
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_materials'))

        file = request.files['file']

        if file.filename == '':
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_materials'))

        if not file.filename.lower().endswith(('.xlsx', '.xls')):
            flash('Пожалуйста, загрузите файл в формате Excel (.xlsx или .xls)', 'danger')
            return redirect(url_for('import_materials'))

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            success, message = materials_db.validate_and_import_from_excel(filepath)
            if success:
                flash(message, 'success')
            else:
                flash(message, 'danger')
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)

        return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в import_materials: {e}")
        traceback.print_exc()
        flash(f'Ошибка при импорте: {str(e)}', 'danger')
        return redirect(url_for('import_materials'))


@app.route('/materials/export')
@manager_required
def export_materials():
    try:
        filepath = materials_db.export_to_excel()
        if filepath and os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name=os.path.basename(filepath))
        else:
            flash('Ошибка при экспорте данных', 'danger')
            return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в export_materials: {e}")
        traceback.print_exc()
        flash(f'Ошибка при экспорте: {str(e)}', 'danger')
        return redirect(url_for('materials_page'))


@app.route('/materials/search', methods=['POST'])
@manager_required
def search_materials():
    try:
        query = request.form.get('query', '').strip()
        search_type = request.form.get('search_type', 'both')
        sort_by = request.form.get('sort_by', 'sap_code')
        sort_order = request.form.get('sort_order', 'asc')

        materials = materials_db.get_all_materials()

        # Фильтрация
        if query:
            filtered = []
            query_lower = query.lower()
            for material in materials:
                if search_type == 'sap':
                    if query in str(material['sap_code']):
                        filtered.append(material)
                elif search_type == 'name':
                    if query_lower in material['material_name'].lower():
                        filtered.append(material)
                else:
                    if query in str(material['sap_code']) or query_lower in material['material_name'].lower():
                        filtered.append(material)
            materials = filtered

        # Сортировка
        reverse = sort_order == 'desc'
        if sort_by == 'sap_code':
            materials.sort(key=lambda x: x['sap_code'], reverse=reverse)
        elif sort_by == 'material_name':
            materials.sort(key=lambda x: x['material_name'].lower(), reverse=reverse)
        elif sort_by == 'palletization':
            materials.sort(key=lambda x: x['palletization'], reverse=reverse)
        elif sort_by == 'box_quantity':
            materials.sort(key=lambda x: x['box_quantity'], reverse=reverse)

        return render_template('materials.html',
                               materials=materials,
                               query=query,
                               search_type=search_type,
                               sort_by=sort_by,
                               sort_order=sort_order)
    except Exception as e:
        print(f"Ошибка в search_materials: {e}")
        traceback.print_exc()
        flash(f'Ошибка при поиске: {str(e)}', 'danger')
        return redirect(url_for('materials_page'))


@app.route('/upload', methods=['POST'])
@labels_access_required
def upload_file():
    try:
        if 'file' not in request.files:
            flash('Файл не выбран', 'warning')
            return redirect(url_for('index'))

        file = request.files['file']

        if file.filename == '':
            flash('Файл не выбран', 'warning')
            return redirect(url_for('index'))

        if not file.filename.lower().endswith('.pdf'):
            flash('Пожалуйста, загрузите файл в формате PDF', 'danger')
            return redirect(url_for('index'))

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            data = extract_data_from_pdf(filepath)
            os.remove(filepath)

            if not data or not data['дата_приемки'] and not data['название_поставщика'] and not data['товары']:
                flash('Не удалось извлечь данные из PDF-файла. Проверьте формат файла.', 'danger')
                return redirect(url_for('index'))

            total_pallets = 0
            for product in data['товары']:
                sap_code = product.get('номенклатурный_номер')
                product['sap_code_display'] = sap_code if sap_code else '-'

                if sap_code:
                    try:
                        sap_code_int = int(sap_code)
                        material = materials_db.get_material_by_sap(sap_code_int)
                        if material:
                            product['palletization'] = material.get('palletization', 0)
                        else:
                            product['palletization'] = 0
                    except (ValueError, TypeError):
                        product['palletization'] = 0
                else:
                    product['palletization'] = 0

                quantity = product.get('количество', 0)
                if quantity == 'Не найдено' or quantity is None:
                    quantity = 0
                else:
                    try:
                        quantity = int(quantity)
                    except:
                        quantity = 0

                palletization = product.get('palletization', 0)
                product['pallets'] = calculate_pallets(quantity, palletization)
                total_pallets += product['pallets']

            db.log_action(session['user'], 'UPLOAD_PDF', f"Загружен PDF: {filename}")

            return render_template('result.html', data=data, total_pallets=total_pallets)
        except Exception as e:
            if os.path.exists(filepath):
                os.remove(filepath)
            print(f"❌ Ошибка при обработке: {e}")
            traceback.print_exc()
            flash(f'Ошибка при обработке файла: {str(e)}', 'danger')
            return redirect(url_for('index'))
    except Exception as e:
        print(f"❌ Ошибка в upload_file: {e}")
        traceback.print_exc()
        flash(f'Ошибка при загрузке файла: {str(e)}', 'danger')
        return redirect(url_for('index'))


@app.route('/recalculate_pallets', methods=['POST'])
@labels_access_required
def recalculate_pallets():
    try:
        data = request.get_json()
        if not data or 'products' not in data:
            return jsonify({'success': False, 'error': 'Нет данных'})

        products = data.get('products', [])
        recalculated = []

        for product in products:
            quantity = product.get('количество', 0)
            palletization = product.get('palletization', 0)
            pallets = calculate_pallets(quantity, palletization)
            recalculated.append(pallets)

        return jsonify({'success': True, 'pallets': recalculated})
    except Exception as e:
        print(f"Ошибка в recalculate_pallets: {e}")
        return jsonify({'success': False, 'error': str(e)})


@app.route('/generate_labels', methods=['POST'])
@labels_access_required
def generate_labels():
    try:
        products_json = request.form.get('products')
        order_number = request.form.get('order_number')
        supplier_name = request.form.get('supplier_name', '')
        order_date = request.form.get('order_date', '')
        selected_indices_json = request.form.get('selected_indices', '[]')

        if not products_json or not order_number:
            flash('Нет данных для формирования этикеток', 'danger')
            return redirect(url_for('index'))

        products_json = products_json.replace("'", '"')

        try:
            products = json.loads(products_json)
        except json.JSONDecodeError as e:
            print(f"Ошибка парсинга JSON: {e}")
            flash(f'Ошибка парсинга данных: {str(e)}', 'danger')
            return redirect(url_for('index'))

        if not products:
            flash('Нет товаров для формирования этикеток', 'danger')
            return redirect(url_for('index'))

        try:
            selected_indices = json.loads(selected_indices_json)
        except json.JSONDecodeError:
            selected_indices = list(range(len(products)))

        selected_indices = [i for i in selected_indices if i < len(products)]

        if not selected_indices:
            flash('Не выбрано ни одного товара для формирования этикеток', 'warning')
            return redirect(url_for('index'))

        settings = load_settings()
        save_path = settings.get('save_path')
        sop_code = settings.get('sop_code', 'RUPD.VLG.05.04.C01')

        filepath, total_labels = create_pallet_labels_file(products, order_number, supplier_name, order_date,
                                                           selected_indices, save_path, sop_code)

        if total_labels == 0:
            flash('Не сгенерировано ни одной этикетки (нет выбранных товаров)', 'warning')
            return redirect(url_for('index'))

        if os.environ.get('RENDER'):
            return render_template('success_render.html',
                                   filepath=filepath,
                                   filename=os.path.basename(filepath),
                                   product_count=len(products),
                                   order_number=order_number,
                                   total_pallets=total_labels)
        else:
            open_in_word(filepath)

        db.log_action(session['user'], 'GENERATE_LABELS',
                      f"Сгенерировано {total_labels} этикеток для заказа {order_number}")

        return render_template('success.html',
                               filepath=filepath,
                               product_count=len(products),
                               order_number=order_number,
                               total_pallets=total_labels)
    except Exception as e:
        print(f"Ошибка в generate_labels: {e}")
        traceback.print_exc()
        flash(f'Ошибка при формировании этикеток: {str(e)}', 'danger')
        return redirect(url_for('index'))


@app.route('/download/<path:filename>')
@login_required
def download_file(filename):
    """Скачивание сгенерированного файла"""
    try:
        settings = load_settings()
        save_path = settings.get('save_path', os.path.join(BASE_DIR, 'generated_labels'))
        filepath = os.path.join(save_path, filename)

        if not os.path.exists(filepath):
            flash('Файл не найден', 'danger')
            return redirect(url_for('index'))

        return send_file(filepath, as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Ошибка при скачивании: {e}")
        traceback.print_exc()
        flash('Ошибка при скачивании файла', 'danger')
        return redirect(url_for('index'))


@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings_page():
    try:
        settings = load_settings()
        drives = get_available_drives()

        current_path = settings.get('save_path', os.path.join(BASE_DIR, 'generated_labels'))

        path_param = request.args.get('path')
        if path_param and os.path.exists(path_param):
            current_path = path_param

        path_parts = []
        if current_path:
            normalized_path = os.path.normpath(current_path)
            parts = normalized_path.split(os.sep)
            temp_path = ""
            for part in parts:
                if part:
                    if temp_path:
                        temp_path = os.path.join(temp_path, part)
                    else:
                        if sys.platform == 'win32' and ':' in part:
                            temp_path = part + os.sep
                        else:
                            temp_path = os.sep + part
                    path_parts.append({
                        'name': part,
                        'path': temp_path
                    })

        subdirs = get_subdirectories(current_path)

        if request.method == 'POST':
            selected_path = request.form.get('selected_path')
            if selected_path:
                settings['save_path'] = selected_path
                save_settings(settings)
                os.makedirs(selected_path, exist_ok=True)
                flash('Настройки сохранены', 'success')
                return redirect(url_for('settings_page'))

            sop_code = request.form.get('sop_code')
            if sop_code:
                settings['sop_code'] = sop_code.strip()
                save_settings(settings)
                flash('Номер СОП успешно обновлен', 'success')
                return redirect(url_for('settings_page'))

            font_settings = {}
            font_keys = ['header_size', 'product_size_short', 'product_size_medium',
                         'product_size_long', 'supplier_size_short', 'supplier_size_medium',
                         'supplier_size_long', 'analytic_size', 'code_size',
                         'line_spacing_short', 'line_spacing_medium', 'line_spacing_long',
                         'short_name_threshold', 'medium_name_threshold']

            for key in font_keys:
                value = request.form.get(key)
                if value is not None:
                    try:
                        if key in ['short_name_threshold', 'medium_name_threshold']:
                            font_settings[key] = int(value)
                        else:
                            font_settings[key] = float(value) if 'line_spacing' in key else int(value)
                    except ValueError:
                        flash(f'Неверное значение для {key}', 'danger')
                        return redirect(url_for('settings_page'))

            if font_settings:
                if 'font_settings' not in settings:
                    settings['font_settings'] = {}
                settings['font_settings'].update(font_settings)
                save_settings(settings)
                flash('Настройки шрифтов успешно сохранены', 'success')
                return redirect(url_for('settings_page'))

            path_component = request.form.get('path_component')
            if path_component:
                new_path = path_component
                settings['save_path'] = new_path
                save_settings(settings)
                os.makedirs(new_path, exist_ok=True)
                flash('Настройки сохранены', 'success')
                return redirect(url_for('settings_page'))

        material_count = materials_db.get_materials_count()
        font_settings = settings.get('font_settings', {})
        is_admin = session.get('user_rights') == 'Администратор'
        is_manager = session.get('user_rights') in ['Администратор', 'Менеджер ОСЛ']

        return render_template('settings.html',
                               settings=settings,
                               drives=drives,
                               current_path=current_path,
                               path_parts=path_parts,
                               subdirs=subdirs,
                               os_sep=os.sep,
                               is_admin=is_admin,
                               is_manager=is_manager,
                               material_count=material_count,
                               font_settings=font_settings)
    except Exception as e:
        print(f"Ошибка в settings_page: {e}")
        traceback.print_exc()
        flash(f'Ошибка загрузки настроек: {str(e)}', 'danger')
        return redirect(url_for('index'))


@app.route('/functionality_not_available')
@login_required
def functionality_not_available():
    flash('Функционал будет доступен в следующей версии', 'info')
    return redirect(url_for('index'))


@app.route('/upload_form')
@labels_access_required
def upload_form():
    try:
        return render_template('upload.html')
    except Exception as e:
        print(f"Ошибка в upload_form: {e}")
        traceback.print_exc()
        flash('Ошибка загрузки страницы', 'danger')
        return redirect(url_for('index'))


@app.route('/exit')
def exit_app():
    return render_template('exit.html')


if __name__ == '__main__':
    load_settings()
    app.run(debug=True, host='0.0.0.0', port=5000)
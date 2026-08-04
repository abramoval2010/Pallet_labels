import sys
import os
import re
import json
import sqlite3
import hashlib
import secrets
import math
from datetime import datetime
from pathlib import Path
import subprocess
import tempfile
import shutil
from contextlib import contextmanager
import io

from flask import Flask, render_template, request, redirect, url_for, send_file, session, flash, jsonify
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)
app.secret_key = secrets.token_hex(32)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['SETTINGS_FILE'] = 'settings.json'
app.config['DATABASE_FILE'] = 'users.db'
app.config['MATERIALS_DATABASE_FILE'] = 'materials.db'

# Создаем папку для загрузок
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# ============================================================
# БАЗА ДАННЫХ МАСТЕР-ДАННЫХ МАТЕРИАЛОВ
# ============================================================

class MaterialsDatabase:
    """Класс для работы с базой мастер-данных материалов"""

    def __init__(self, db_file='materials.db'):
        self.db_file = db_file
        self.init_database()

    @contextmanager
    def get_connection(self):
        """Получает соединение с базой данных"""
        conn = sqlite3.connect(self.db_file)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
        finally:
            conn.close()

    def init_database(self):
        """Инициализирует базу данных и создает таблицу материалов"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS materials (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    sap_code INTEGER UNIQUE NOT NULL,
                    material_name TEXT UNIQUE NOT NULL,
                    palletization INTEGER DEFAULT 0,
                    box_quantity INTEGER DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            conn.commit()

            cursor.execute("SELECT COUNT(*) FROM materials")
            count = cursor.fetchone()[0]

            if count == 0:
                self.create_default_materials()

    def create_default_materials(self):
        """Создает тестовые материалы"""
        default_materials = [
            {'sap_code': 10000001, 'material_name': 'Тестовый материал 1', 'palletization': 100, 'box_quantity': 10},
            {'sap_code': 10000002, 'material_name': 'Тестовый материал 2', 'palletization': 200, 'box_quantity': 20},
            {'sap_code': 10000003, 'material_name': 'Тестовый материал 3', 'palletization': 150, 'box_quantity': 15},
        ]

        for material in default_materials:
            self.add_material(material)

        print(f"✅ Создано {len(default_materials)} тестовых материалов")

    def add_material(self, material_data):
        """Добавляет новый материал в базу"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO materials (sap_code, material_name, palletization, box_quantity)
                    VALUES (?, ?, ?, ?)
                ''', (
                    material_data['sap_code'],
                    material_data['material_name'],
                    material_data.get('palletization', 0),
                    material_data.get('box_quantity', 0)
                ))
                conn.commit()
                return True, "Материал успешно добавлен"
        except sqlite3.IntegrityError as e:
            if 'UNIQUE constraint failed: materials.sap_code' in str(e):
                return False, f"Материал с SAP-кодом {material_data['sap_code']} уже существует"
            elif 'UNIQUE constraint failed: materials.material_name' in str(e):
                return False, f"Материал с названием '{material_data['material_name']}' уже существует"
            return False, f"Ошибка целостности данных: {str(e)}"
        except Exception as e:
            return False, f"Ошибка при добавлении материала: {str(e)}"

    def get_material_by_sap(self, sap_code):
        """Находит материал по SAP-коду"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM materials WHERE sap_code = ?", (sap_code,))
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None

    def get_material_by_name(self, material_name):
        """Находит материал по названию"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM materials WHERE material_name = ?", (material_name,))
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None

    def update_material(self, sap_code, new_data):
        """Обновляет данные материала"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()

                fields = []
                values = []

                for key in ['material_name', 'palletization', 'box_quantity']:
                    if key in new_data:
                        fields.append(f"{key} = ?")
                        values.append(new_data[key])

                values.append(sap_code)
                fields.append("updated_at = CURRENT_TIMESTAMP")

                query = f"UPDATE materials SET {', '.join(fields)} WHERE sap_code = ?"
                cursor.execute(query, values)
                conn.commit()
                return True, "Данные материала обновлены"
        except sqlite3.IntegrityError as e:
            if 'UNIQUE constraint failed: materials.material_name' in str(e):
                return False, f"Материал с названием '{new_data.get('material_name', '')}' уже существует"
            return False, f"Ошибка целостности данных: {str(e)}"
        except Exception as e:
            return False, f"Ошибка при обновлении: {str(e)}"

    def delete_material(self, sap_code):
        """Удаляет материал"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("DELETE FROM materials WHERE sap_code = ?", (sap_code,))
                conn.commit()
                return True, "Материал удален"
        except Exception as e:
            return False, f"Ошибка при удалении: {str(e)}"

    def get_all_materials(self):
        """Возвращает список всех материалов"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM materials ORDER BY sap_code")
            rows = cursor.fetchall()
            return [dict(row) for row in rows]

    def get_materials_count(self):
        """Возвращает количество материалов в базе"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM materials")
            return cursor.fetchone()[0]

    def validate_and_import_from_excel(self, filepath):
        """Проверяет и импортирует данные из Excel файла"""
        try:
            df = pd.read_excel(filepath, header=None)

            if df.empty:
                return False, "Файл пуст"

            header_row = 0
            first_row = df.iloc[0]
            is_header = False

            if isinstance(first_row[0], str) and not first_row[0].isdigit():
                is_header = True
            elif isinstance(first_row[1], str) and len(str(first_row[1])) > 0:
                is_header = True

            if is_header:
                header_row = 1
                df = df.iloc[1:].reset_index(drop=True)

            if df.shape[1] < 3:
                return False, "Файл должен содержать минимум 3 колонки (SAP код, Название, Палетизация)"

            errors = []
            materials_to_add = []
            materials_to_update = []

            for idx, row in df.iterrows():
                sap_value = row[0]
                if pd.isna(sap_value):
                    errors.append(f"Строка {idx + 2}: SAP код не может быть пустым")
                    continue

                try:
                    sap_code = int(sap_value)
                    if len(str(sap_code)) != 8:
                        errors.append(
                            f"Строка {idx + 2}: SAP код должен содержать 8 цифр (текущее значение: {sap_code})")
                        continue
                except (ValueError, TypeError):
                    errors.append(f"Строка {idx + 2}: SAP код должен быть числом (текущее значение: {sap_value})")
                    continue

                material_name = str(row[1]) if not pd.isna(row[1]) else ''
                if not material_name.strip():
                    errors.append(f"Строка {idx + 2}: Название материала не может быть пустым")
                    continue

                palletization_value = row[2] if not pd.isna(row[2]) else 0
                try:
                    palletization = int(palletization_value)
                    if palletization < 0:
                        errors.append(
                            f"Строка {idx + 2}: Палетизация не может быть отрицательной (текущее значение: {palletization})")
                        continue
                except (ValueError, TypeError):
                    errors.append(
                        f"Строка {idx + 2}: Палетизация должна быть числом (текущее значение: {palletization_value})")
                    continue

                if df.shape[1] > 3:
                    for col_idx in range(3, df.shape[1]):
                        if not pd.isna(row[col_idx]) and str(row[col_idx]).strip():
                            errors.append(f"Строка {idx + 2}: Обнаружены данные в дополнительной колонке {col_idx + 1}")
                            break

                material_data = {
                    'sap_code': sap_code,
                    'material_name': material_name.strip(),
                    'palletization': palletization
                }

                existing = self.get_material_by_sap(sap_code)
                if existing:
                    materials_to_update.append(material_data)
                else:
                    existing_name = self.get_material_by_name(material_name.strip())
                    if existing_name:
                        errors.append(
                            f"Строка {idx + 2}: Материал с названием '{material_name.strip()}' уже существует в базе (SAP: {existing_name['sap_code']})")
                    else:
                        materials_to_add.append(material_data)

            if errors:
                return False, "Ошибки в файле:\n" + "\n".join(errors)

            updated_count = 0
            for material in materials_to_update:
                existing = self.get_material_by_sap(material['sap_code'])
                if existing:
                    new_data = {
                        'material_name': material['material_name'],
                        'palletization': material['palletization'],
                        'box_quantity': existing.get('box_quantity', 0)
                    }
                    success, message = self.update_material(material['sap_code'], new_data)
                    if success:
                        updated_count += 1

            added_count = 0
            for material in materials_to_add:
                material_data = {
                    'sap_code': material['sap_code'],
                    'material_name': material['material_name'],
                    'palletization': material['palletization'],
                    'box_quantity': 0
                }
                success, message = self.add_material(material_data)
                if success:
                    added_count += 1

            return True, f"Успешно обработано: добавлено {added_count} материалов, обновлено {updated_count} материалов"

        except Exception as e:
            return False, f"Ошибка при обработке файла: {str(e)}"


# ============================================================
# ШИФРОВАННАЯ БАЗА ДАННЫХ ПОЛЬЗОВАТЕЛЕЙ (SQLite)
# ============================================================

class UserDatabase:
    """Класс для работы с зашифрованной базой данных пользователей"""

    def __init__(self, db_file='users.db'):
        self.db_file = db_file
        self.master_key = self._get_master_key()
        self.init_database()

    def _get_master_key(self):
        """Получает или создает мастер-ключ для шифрования"""
        key_file = 'master.key'
        if os.path.exists(key_file):
            with open(key_file, 'rb') as f:
                return f.read()
        else:
            key = secrets.token_bytes(32)
            with open(key_file, 'wb') as f:
                f.write(key)
            return key

    def _encrypt(self, data):
        """Шифрует данные"""
        if not data:
            return ''
        key = self.master_key
        encrypted = bytearray()
        for i, char in enumerate(str(data).encode('utf-8')):
            encrypted.append(char ^ key[i % len(key)])
        return encrypted.hex()

    def _decrypt(self, encrypted_data):
        """Расшифровывает данные"""
        if not encrypted_data:
            return ''
        key = self.master_key
        encrypted_bytes = bytes.fromhex(encrypted_data)
        decrypted = bytearray()
        for i, byte in enumerate(encrypted_bytes):
            decrypted.append(byte ^ key[i % len(key)])
        return decrypted.decode('utf-8')

    @contextmanager
    def get_connection(self):
        """Получает соединение с базой данных"""
        conn = sqlite3.connect(self.db_file)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
        finally:
            conn.close()

    def init_database(self):
        """Инициализирует базу данных и создает таблицы"""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    login TEXT UNIQUE NOT NULL,
                    email TEXT NOT NULL,
                    password_hash TEXT NOT NULL,
                    fname TEXT,
                    lname TEXT,
                    enterprise TEXT,
                    plant TEXT,
                    rights TEXT,
                    status TEXT DEFAULT 'active',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            cursor.execute("PRAGMA table_info(users)")
            columns = [col[1] for col in cursor.fetchall()]
            if 'status' not in columns:
                print("🔧 Добавляем колонку status в таблицу users")
                cursor.execute("ALTER TABLE users ADD COLUMN status TEXT DEFAULT 'active'")
                conn.commit()
                cursor.execute("UPDATE users SET status = 'active' WHERE status IS NULL")
                conn.commit()

            cursor.execute('''
                CREATE TABLE IF NOT EXISTS audit_log (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_login TEXT,
                    action TEXT,
                    details TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            conn.commit()

            cursor.execute("SELECT COUNT(*) FROM users")
            count = cursor.fetchone()[0]

            if count == 0:
                self.create_default_users()

    def create_default_users(self):
        """Создает пользователей по умолчанию"""
        default_users = [
            {
                'login': 'admin',
                'email': 'admin@example.com',
                'password': 'admin123',
                'fname': 'Администратор',
                'lname': 'Системы',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Администратор',
                'status': 'active'
            },
            {
                'login': 'operator',
                'email': 'operator@example.com',
                'password': 'operator123',
                'fname': 'Оператор',
                'lname': 'СиМ',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Оператор СиМ',
                'status': 'active'
            },
            {
                'login': 'alexander.abramov@abbott.com',
                'email': 'alexander.abramov@abbott.com',
                'password': 'alexander.abramov@abbott.com',
                'fname': 'Александр',
                'lname': 'Абрамов',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Администратор',
                'status': 'active'
            },
            {
                'login': 'anastasia.yudashkina@abbott.com',
                'email': 'anastasia.yudashkina@abbott.com',
                'password': 'anastasia.yudashkina@abbott.com',
                'fname': 'Анастасия',
                'lname': 'Юдашкина',
                'enterprise': 'ООО "Верофарм"',
                'plant': 'Вольгинский',
                'rights': 'Оператор СиМ',
                'status': 'active'
            }
        ]

        for user in default_users:
            self.add_user(user)

        print(f"✅ Создано {len(default_users)} пользователей в базе данных")

    def add_user(self, user_data):
        """Добавляет нового пользователя в базу"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                password_hash = generate_password_hash(user_data['password'])

                cursor.execute('''
                    INSERT INTO users (login, email, password_hash, fname, lname, enterprise, plant, rights, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    user_data['login'],
                    user_data['email'],
                    password_hash,
                    user_data.get('fname', ''),
                    user_data.get('lname', ''),
                    user_data.get('enterprise', ''),
                    user_data.get('plant', ''),
                    user_data.get('rights', 'Оператор СиМ'),
                    user_data.get('status', 'active')
                ))

                conn.commit()
                self.log_action(user_data['login'], 'ADD_USER', f"Добавлен пользователь {user_data['login']}")
                return True, "Пользователь успешно добавлен"
        except sqlite3.IntegrityError:
            return False, "Пользователь с таким логином уже существует"
        except Exception as e:
            return False, f"Ошибка при добавлении пользователя: {str(e)}"

    def find_user(self, login):
        """Находит пользователя по логину"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM users WHERE login = ?", (login,))
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None

    def authenticate(self, login, password):
        """Проверяет логин и пароль"""
        user = self.find_user(login)
        if user:
            if user.get('status') == 'blocked':
                self.log_action(login, 'LOGIN_FAILED', f"Попытка входа заблокированного пользователя")
                return None, 'blocked'
            if user.get('status') == 'deleted':
                self.log_action(login, 'LOGIN_FAILED', f"Попытка входа удаленного пользователя")
                return None, 'deleted'

            if check_password_hash(user['password_hash'], password):
                self.log_action(login, 'LOGIN_SUCCESS', f"Успешный вход")
                return user, 'success'
            else:
                self.log_action(login, 'LOGIN_FAILED', f"Неверный пароль")
        return None, 'invalid'

    def update_password(self, login, new_password):
        """Обновляет пароль пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                password_hash = generate_password_hash(new_password)
                cursor.execute(
                    "UPDATE users SET password_hash = ?, updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (password_hash, login)
                )
                conn.commit()
                self.log_action(login, 'UPDATE_PASSWORD', "Пароль изменен")
                return True
        except Exception as e:
            print(f"Ошибка обновления пароля: {e}")
            return False

    def update_user(self, login, new_data):
        """Обновляет данные пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()

                fields = []
                values = []

                for key in ['email', 'fname', 'lname', 'enterprise', 'plant', 'rights', 'status']:
                    if key in new_data:
                        fields.append(f"{key} = ?")
                        values.append(new_data[key])

                if 'password' in new_data and new_data['password']:
                    fields.append("password_hash = ?")
                    values.append(generate_password_hash(new_data['password']))

                values.append(login)
                fields.append("updated_at = CURRENT_TIMESTAMP")

                query = f"UPDATE users SET {', '.join(fields)} WHERE login = ?"
                cursor.execute(query, values)
                conn.commit()

                self.log_action(login, 'UPDATE_USER', f"Данные пользователя обновлены")
                return True, "Данные пользователя обновлены"
        except Exception as e:
            return False, f"Ошибка обновления: {str(e)}"

    def delete_user(self, login):
        """Удаляет пользователя (soft delete - меняет статус)"""
        admins = self.get_users_by_rights('Администратор')
        user = self.find_user(login)

        if user and user['rights'] == 'Администратор' and len(admins) <= 1:
            return False, "Нельзя удалить единственного администратора"

        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE users SET status = 'deleted', updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (login,)
                )
                conn.commit()
                self.log_action(login, 'DELETE_USER', f"Пользователь удален")
                return True, "Пользователь удален"
        except Exception as e:
            return False, f"Ошибка удаления: {str(e)}"

    def block_user(self, login):
        """Блокирует пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE users SET status = 'blocked', updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (login,)
                )
                conn.commit()
                self.log_action(login, 'BLOCK_USER', f"Пользователь заблокирован")
                return True, "Пользователь заблокирован"
        except Exception as e:
            return False, f"Ошибка блокировки: {str(e)}"

    def unblock_user(self, login):
        """Разблокирует пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE users SET status = 'active', updated_at = CURRENT_TIMESTAMP WHERE login = ?",
                    (login,)
                )
                conn.commit()
                self.log_action(login, 'UNBLOCK_USER', f"Пользователь разблокирован")
                return True, "Пользователь разблокирован"
        except Exception as e:
            return False, f"Ошибка разблокировки: {str(e)}"

    def get_all_users(self):
        """Возвращает список всех пользователей"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id, login, email, fname, lname, enterprise, plant, rights, status, created_at FROM users ORDER BY login")
            rows = cursor.fetchall()
            return [dict(row) for row in rows]

    def get_users_by_rights(self, rights):
        """Возвращает пользователей с определенными правами"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM users WHERE rights = ? AND status = 'active'", (rights,))
            rows = cursor.fetchall()
            return [dict(row) for row in rows]

    def log_action(self, user_login, action, details):
        """Логирует действие пользователя"""
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO audit_log (user_login, action, details)
                    VALUES (?, ?, ?)
                ''', (user_login, action, details))
                conn.commit()
        except Exception as e:
            print(f"Ошибка логирования: {e}")

    def get_audit_log(self, limit=100):
        """Возвращает последние записи аудита"""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT * FROM audit_log ORDER BY timestamp DESC LIMIT ?",
                (limit,)
            )
            rows = cursor.fetchall()
            return [dict(row) for row in rows]


# ============================================================
# ИНИЦИАЛИЗАЦИЯ БАЗ ДАННЫХ
# ============================================================

db = UserDatabase(app.config['DATABASE_FILE'])
materials_db = MaterialsDatabase(app.config['MATERIALS_DATABASE_FILE'])

# ============================================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С НАСТРОЙКАМИ
# ============================================================

SETTINGS_FILE = 'settings.json'


def load_settings():
    """Загружает настройки из файла"""
    default_settings = {
        'save_path': os.path.join(os.path.expanduser('~'), 'Documents', 'Палетные этикетки'),
        'sop_code': 'RUPD.VLG.05.04.C01'
    }

    try:
        if os.path.exists(SETTINGS_FILE):
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                settings = json.load(f)
                if 'save_path' not in settings:
                    settings['save_path'] = default_settings['save_path']
                if 'sop_code' not in settings:
                    settings['sop_code'] = default_settings['sop_code']
                return settings
        else:
            os.makedirs(default_settings['save_path'], exist_ok=True)
            save_settings(default_settings)
            return default_settings
    except Exception as e:
        print(f"Ошибка загрузки настроек: {e}")
        return default_settings


def save_settings(settings):
    """Сохраняет настройки в файл"""
    try:
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
        return True
    except Exception as e:
        print(f"Ошибка сохранения настроек: {e}")
        return False


def get_available_drives():
    """Получает список доступных дисков в Windows"""
    drives = []
    if sys.platform == 'win32':
        import string
        for letter in string.ascii_uppercase:
            drive = f"{letter}:\\"
            if os.path.exists(drive):
                drives.append(drive)
    else:
        drives = ['/']
    return drives


def get_subdirectories(path):
    """Получает список подпапок в указанном пути"""
    subdirs = []
    try:
        if os.path.exists(path):
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                if os.path.isdir(item_path):
                    subdirs.append(item)
    except:
        pass
    return sorted(subdirs)


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================

def parse_quantity_from_string(text):
    """
    Извлекает число из строки, учитывая пробелы между разрядами.
    Например: "24 000" -> 24000, "1 500" -> 1500
    """
    if not text:
        return None

    # Удаляем все пробелы из строки
    cleaned = text.replace(' ', '').replace('\u00a0', '').replace('\t', '')

    # Пробуем найти число в строке
    match = re.search(r'(\d+)', cleaned)
    if match:
        try:
            return int(match.group(1))
        except:
            return None
    return None


def calculate_pallets(quantity, palletization):
    """
    Рассчитывает количество паллет по формуле:
    - если палетизация > 0: ceil(количество / палетизация)
    - если палетизация <= 0: 1
    """
    if not quantity:
        return 1

    try:
        qty = int(quantity)
    except:
        return 1

    if palletization > 0:
        try:
            return math.ceil(qty / palletization)
        except:
            return 1
    else:
        return 1


# ============================================================
# ФУНКЦИЯ ГЕНЕРАЦИИ ЭТИКЕТОК
# ============================================================

def create_pallet_labels_file(all_products, order_number, save_path=None, sop_code=None):
    """
    Создает файл с палетными этикетками.
    Каждый аналитический лист дублируется столько раз, сколько указано в поле 'pallets'.
    Каждая этикетка находится на отдельной странице.
    """
    if save_path is None:
        settings = load_settings()
        save_path = settings.get('save_path', os.path.join(os.path.expanduser('~'), 'Documents', 'Палетные этикетки'))

    if sop_code is None:
        settings = load_settings()
        sop_code = settings.get('sop_code', 'RUPD.VLG.05.04.C01')

    try:
        os.makedirs(save_path, exist_ok=True)
    except PermissionError:
        save_path = os.path.join(os.path.expanduser('~'), 'Documents', 'Палетные этикетки')
        os.makedirs(save_path, exist_ok=True)
    except Exception as e:
        print(f"Ошибка создания папки: {e}")
        save_path = os.path.join(os.path.expanduser('~'), 'Documents')

    # Создаем новый документ
    doc = Document()

    # Настраиваем первую страницу
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.15)
        section.bottom_margin = Inches(0.15)
        section.left_margin = Inches(0.15)
        section.right_margin = Inches(0.15)

    total_labels = 0
    first_page = True

    for product in all_products:
        # Получаем количество паллет для этого товара
        pallets_count = product.get('pallets', 1)
        if pallets_count < 1:
            pallets_count = 1

        # Рассчитываем размеры шрифтов для этого товара
        product_name = product.get('наименование', 'Не указано')
        supplier_series = product.get('партия_поставщика', 'Не указана')
        analytic_list = product.get('аналитический_лист', 'Не указан')

        name_len = len(product_name)

        if name_len < 30:
            font_sizes = {'header': 14, 'product': 82, 'supplier': 60, 'analytic': 76, 'code': 14}
        elif name_len < 40:
            font_sizes = {'header': 14, 'product': 74, 'supplier': 58, 'analytic': 68, 'code': 14}
        else:
            font_sizes = {'header': 14, 'product': 70, 'supplier': 56, 'analytic': 66, 'code': 14}

        if name_len > 55:
            font_sizes['product'] = 62
            font_sizes['supplier'] = 50
            font_sizes['analytic'] = 58
        elif name_len > 70:
            font_sizes['product'] = 54
            font_sizes['supplier'] = 44
            font_sizes['analytic'] = 50
        elif name_len > 85:
            font_sizes['product'] = 46
            font_sizes['supplier'] = 38
            font_sizes['analytic'] = 44

        supplier_len = len(supplier_series) if supplier_series else 0
        if supplier_len > 25:
            font_sizes['supplier'] = max(32, font_sizes['supplier'] - 12)
        elif supplier_len > 18:
            font_sizes['supplier'] = max(38, font_sizes['supplier'] - 8)

        analytic_len = len(str(analytic_list)) if analytic_list else 0
        if analytic_len > 15:
            font_sizes['analytic'] = max(42, font_sizes['analytic'] - 10)

        # Создаем паллеты для этого товара
        for pallet_num in range(pallets_count):
            # Добавляем разрыв страницы перед каждой этикеткой, кроме первой
            if not first_page:
                doc.add_page_break()
                # Настраиваем новую секцию
                if doc.sections:
                    section = doc.sections[-1]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11.69)
                    section.page_height = Inches(8.27)
                    section.top_margin = Inches(0.15)
                    section.bottom_margin = Inches(0.15)
                    section.left_margin = Inches(0.15)
                    section.right_margin = Inches(0.15)
            else:
                first_page = False

            # Создаем таблицу для этикетки
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(10.8)

            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].clear()

            # ООО «Верофарм»
            p1 = cell.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(2)
            p1.paragraph_format.space_after = Pt(2)
            run1 = p1.add_run('ООО «Верофарм»')
            run1.italic = True
            run1.bold = True
            run1.font.size = Pt(font_sizes['header'])
            run1.font.name = 'Calibri'

            # Название товара
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(5)
            p2.paragraph_format.space_after = Pt(5)
            run2 = p2.add_run(product_name)
            run2.italic = True
            run2.bold = True
            run2.font.size = Pt(font_sizes['product'])
            run2.font.name = 'Calibri'

            # Серия поставщика
            supplier_text = f"п. {supplier_series}" if supplier_series and supplier_series != 'Не указана' else supplier_series
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(4)
            p3.paragraph_format.space_after = Pt(4)
            run3 = p3.add_run(supplier_text)
            run3.italic = True
            run3.bold = False
            run3.font.size = Pt(font_sizes['supplier'])
            run3.font.name = 'Calibri'

            # Аналитический лист
            p4 = cell.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(4)
            p4.paragraph_format.space_after = Pt(4)
            run4 = p4.add_run(f'АЛ {analytic_list}')
            run4.italic = True
            run4.bold = True
            run4.font.size = Pt(font_sizes['analytic'])
            run4.font.name = 'Calibri'

            # Код
            p5 = cell.add_paragraph()
            p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p5.paragraph_format.space_before = Pt(8)
            p5.paragraph_format.space_after = Pt(2)
            run5 = p5.add_run(sop_code)
            run5.italic = False
            run5.bold = True
            run5.font.size = Pt(font_sizes['code'])
            run5.font.name = 'Calibri'

            total_labels += 1

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Информационный лист {order_number}_{timestamp}.docx"
    filepath = os.path.join(save_path, filename)

    try:
        doc.save(filepath)
    except PermissionError:
        fallback_path = os.path.join(os.path.expanduser('~'), 'Documents')
        filepath = os.path.join(fallback_path, filename)
        doc.save(filepath)
    except Exception as e:
        print(f"Ошибка сохранения файла: {e}")
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)

    return filepath


def open_in_word(filepath):
    """Открывает файл в Microsoft Word"""
    try:
        if sys.platform == 'win32':
            os.startfile(filepath)
        elif sys.platform == 'darwin':
            subprocess.run(['open', filepath])
        else:
            subprocess.run(['xdg-open', filepath])
        return True
    except Exception as e:
        print(f"Ошибка при открытии файла: {e}")
        return False


# ============================================================
# ПАРСИНГ PDF
# ============================================================

def extract_data_from_pdf(pdf_path):
    """Извлекает данные из PDF-файла приходного ордера"""
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"Ошибка открытия PDF: {e}")
        return None

    lines = [line.strip() for line in text.split('\n') if line.strip()]

    result = {
        'дата_приемки': None,
        'название_поставщика': None,
        'номер_ордера': None,
        'товары': []
    }

    for line in lines:
        if 'ПРИХОДНЫЙ ОРДЕР №' in line:
            order_match = re.search(r'№\s*(\d+)', line)
            if order_match:
                result['номер_ордера'] = order_match.group(1)
                break

    for line in lines:
        if re.match(r'\d{2}\.\d{2}\.\d{4}', line):
            if '2028' not in line and 'годности' not in line:
                result['дата_приемки'] = line
                break

    if not result['дата_приемки']:
        for i, line in enumerate(lines):
            if 'составления' in line.lower() and i + 1 < len(lines):
                date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', lines[i + 1])
                if date_match:
                    result['дата_приемки'] = date_match.group(1)
                    break

    for line in lines:
        if 'ООО' in line or 'ЗАО' in line or 'ОАО' in line or 'АО' in line or 'ИП' in line:
            if 'ВЕРОФАРМ' not in line and 'Завод' not in line:
                result['название_поставщика'] = line
                break

    start_index = -1
    end_index = -1

    for i, line in enumerate(lines):
        if 'Материальные ценности' in line:
            start_index = i
            break

    for i, line in enumerate(lines):
        if line == 'Итого' or line.startswith('Итого'):
            end_index = i
            break

    if start_index != -1 and end_index != -1:
        data_start = -1
        for i in range(start_index, end_index):
            if lines[i] == '14':
                data_start = i + 1
                break

        if data_start != -1:
            header_keywords = [
                'наименование', 'сорт', 'размер', 'марка', 'код',
                'единица измерения', 'количество', 'цена', 'сумма',
                'без учета', 'всего с учетом', 'номер паспорта',
                'партия', 'поставщика', 'партия поставщика',
                'срок годности', 'паспорта'
            ]

            i = data_start
            products = []

            while i < end_index:
                line = lines[i]

                if not line:
                    i += 1
                    continue

                if re.match(r'\d{2}\.\d{2}\.\d{4}', line):
                    i += 1
                    continue

                if re.match(r'^[\d\s,]+$', line) and len(line) > 3:
                    i += 1
                    continue

                if line == 'X':
                    i += 1
                    continue

                is_header = False
                line_lower = line.lower()
                for keyword in header_keywords:
                    if keyword in line_lower:
                        is_header = True
                        break

                if is_header:
                    i += 1
                    continue

                if re.search(r'[А-ЯЁA-Z]', line) and not re.match(r'^[\d\s,]+$', line):
                    product_name_lines = [line]
                    current_pos = i + 1
                    last_name_line = i

                    # Сохраняем позицию начала названия для последующего извлечения количества
                    name_start_line = i

                    while current_pos < end_index:
                        next_line = lines[current_pos]

                        nomencl_match = re.search(r'\b(\d{8})\b', next_line)
                        if nomencl_match:
                            if not nomencl_match.group(1).startswith(('8', '7')):
                                if nomencl_match.group(1) not in ['48797491', '21000101']:
                                    last_name_line = current_pos - 1
                                    break

                        if re.search(r'[А-ЯЁA-Z]', next_line) and not re.match(r'^[\d\s,]+$', next_line):
                            skip_words = ['кг', 'шт', 'г', 'л', 'м', 'уп', 'пач', 'итого',
                                          'x', 'принял', 'сдал', 'должность', 'подпись',
                                          'расшифровка', 'номер', 'паспорта']
                            if next_line.lower() not in skip_words:
                                product_name_lines.append(next_line)
                                current_pos += 1
                            else:
                                break
                        else:
                            break

                    if last_name_line == i:
                        last_name_line = current_pos - 1

                    product_name = ' '.join(product_name_lines)

                    nomencl_number = None
                    nomencl_line = last_name_line + 1
                    if nomencl_line < len(lines):
                        nomencl_match = re.search(r'\b(\d{8})\b', lines[nomencl_line])
                        if nomencl_match:
                            if not nomencl_match.group(1).startswith(('8', '7')):
                                if nomencl_match.group(1) not in ['48797491', '21000101']:
                                    nomencl_number = nomencl_match.group(1)

                    analytic_list = None
                    batch_line = last_name_line + 8
                    if batch_line < len(lines):
                        batch_match = re.search(r'\b(1\d{9})\b', lines[batch_line])
                        if batch_match:
                            analytic_list = batch_match.group(1)

                    supplier_batch = None
                    supplier_line = last_name_line + 9
                    if supplier_line < len(lines):
                        supplier_batch = lines[supplier_line]

                    expiration_date = None
                    expiration_line = last_name_line + 10
                    if expiration_line < len(lines):
                        date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', lines[expiration_line])
                        if date_match:
                            expiration_date = date_match.group(1)

                    # Извлекаем количество из строки на 5 позиций ниже от начала названия
                    quantity = None
                    quantity_line_index = name_start_line + 5
                    if quantity_line_index < len(lines):
                        quantity_line = lines[quantity_line_index]
                        quantity = parse_quantity_from_string(quantity_line)

                    if analytic_list:
                        product_data = {
                            'наименование': product_name,
                            'номенклатурный_номер': nomencl_number,
                            'аналитический_лист': analytic_list,
                            'партия_поставщика': supplier_batch,
                            'срок_годности': expiration_date,
                            'количество': quantity
                        }
                        products.append(product_data)

                    i = last_name_line + 11
                else:
                    i += 1

            result['товары'] = products

    return result


# ============================================================
# ДЕКОРАТОРЫ ДЛЯ ПРОВЕРКИ ПРАВ ДОСТУПА
# ============================================================

def login_required(f):
    """Декоратор для проверки авторизации"""

    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


def admin_required(f):
    """Декоратор для проверки прав администратора"""

    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))
        if session.get('user_rights') != 'Администратор':
            flash('Доступ запрещен. Требуются права администратора.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


def labels_access_required(f):
    """Декоратор для проверки доступа к формированию этикеток"""

    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Пожалуйста, войдите в систему', 'warning')
            return redirect(url_for('login'))

        user_rights = session.get('user_rights')
        user_plant = session.get('user_plant')

        has_access = False

        if user_rights == 'Администратор':
            has_access = True
        elif user_rights == 'Оператор СиМ' and user_plant == 'Вольгинский':
            has_access = True
        elif user_rights == 'Менеджер ОСЛ' and user_plant == 'Вольгинский':
            has_access = True

        if not has_access:
            flash('Доступ запрещен. У вас недостаточно прав для формирования этикеток.', 'danger')
            return redirect(url_for('index'))

        return f(*args, **kwargs)

    decorated_function.__name__ = f.__name__
    return decorated_function


# ============================================================
# МАРШРУТЫ ПРИЛОЖЕНИЯ
# ============================================================

@app.route('/')
def index():
    """Главная страница - после входа показывает главный экран"""
    try:
        if 'user' in session:
            settings = load_settings()
            return render_template('index.html', settings=settings)
        return redirect(url_for('login'))
    except Exception as e:
        print(f"Ошибка в index: {e}")
        return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    """Страница входа в систему"""
    try:
        if request.method == 'POST':
            login = request.form.get('login', '').strip()
            password = request.form.get('password', '')
            remember = request.form.get('remember', False)

            if not login or not password:
                flash('Введите логин и пароль', 'warning')
                return render_template('login.html')

            user, status = db.authenticate(login, password)

            if status == 'blocked':
                flash('Пользователь заблокирован. Обратитесь к администратору.', 'danger')
                return render_template('login.html')

            if status == 'deleted':
                flash('Пользователь удален. Обратитесь к администратору.', 'danger')
                return render_template('login.html')

            if user:
                session['user'] = user['login']
                session['user_email'] = user['email']
                session['user_fname'] = user['fname']
                session['user_lname'] = user['lname']
                session['user_enterprise'] = user['enterprise']
                session['user_plant'] = user['plant']
                session['user_rights'] = user['rights']

                if remember:
                    session.permanent = True

                flash(f'Добро пожаловать, {user["fname"]} {user["lname"]}!', 'success')
                return redirect(url_for('index'))
            else:
                flash('Неверный логин или пароль. Проверьте правильность ввода.', 'danger')

        return render_template('login.html')
    except Exception as e:
        print(f"Ошибка в login: {e}")
        flash('Ошибка при входе в систему', 'danger')
        return render_template('login.html')


@app.route('/logout')
def logout():
    """Выход из системы"""
    if 'user' in session:
        db.log_action(session['user'], 'LOGOUT', "Выход из системы")
    session.clear()
    flash('Вы вышли из системы', 'info')
    return redirect(url_for('login'))


@app.route('/change_password', methods=['GET', 'POST'])
@login_required
def change_password():
    """Смена пароля"""
    try:
        if request.method == 'POST':
            current_password = request.form.get('current_password', '')
            new_password = request.form.get('new_password', '')
            confirm_password = request.form.get('confirm_password', '')

            user = db.find_user(session['user'])

            if not user:
                flash('Пользователь не найден', 'danger')
                return redirect(url_for('index'))

            if not check_password_hash(user['password_hash'], current_password):
                flash('Неверный текущий пароль', 'danger')
                return render_template('change_password.html')

            if len(new_password) < 6:
                flash('Новый пароль должен содержать минимум 6 символов', 'warning')
                return render_template('change_password.html')

            if new_password != confirm_password:
                flash('Пароли не совпадают', 'warning')
                return render_template('change_password.html')

            if db.update_password(session['user'], new_password):
                flash('Пароль успешно изменен', 'success')
                return redirect(url_for('index'))
            else:
                flash('Ошибка при смене пароля', 'danger')

        return render_template('change_password.html')
    except Exception as e:
        print(f"Ошибка в change_password: {e}")
        flash('Ошибка при смене пароля', 'danger')
        return redirect(url_for('index'))


@app.route('/admin')
@admin_required
def admin_panel():
    """Панель администратора - управление пользователями"""
    try:
        users = db.get_all_users()
        audit_log = db.get_audit_log(50)
        return render_template('admin.html', users=users, audit_log=audit_log)
    except Exception as e:
        print(f"Ошибка в admin_panel: {e}")
        import traceback
        traceback.print_exc()
        flash(f'Ошибка загрузки панели администратора: {str(e)}', 'danger')
        return redirect(url_for('settings_page'))


@app.route('/admin/add_user', methods=['POST'])
@admin_required
def add_user():
    """Добавление нового пользователя"""
    try:
        login = request.form.get('login', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        fname = request.form.get('fname', '').strip()
        lname = request.form.get('lname', '').strip()
        enterprise = request.form.get('enterprise', 'ООО "Верофарм"')
        plant = request.form.get('plant', 'Вольгинский')
        rights = request.form.get('rights', 'Оператор СиМ')
        status = request.form.get('status', 'active')

        if not login or not email or not password:
            flash('Логин, email и пароль обязательны для заполнения', 'warning')
            return redirect(url_for('admin_panel'))

        if len(password) < 6:
            flash('Пароль должен содержать минимум 6 символов', 'warning')
            return redirect(url_for('admin_panel'))

        user_data = {
            'login': login,
            'email': email,
            'password': password,
            'fname': fname,
            'lname': lname,
            'enterprise': enterprise,
            'plant': plant,
            'rights': rights,
            'status': status
        }

        success, message = db.add_user(user_data)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в add_user: {e}")
        flash('Ошибка при добавлении пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/delete_user/<login>', methods=['POST'])
@admin_required
def delete_user(login):
    """Удаление пользователя (soft delete)"""
    try:
        if login == session['user']:
            flash('Нельзя удалить самого себя', 'danger')
            return redirect(url_for('admin_panel'))

        success, message = db.delete_user(login)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в delete_user: {e}")
        flash('Ошибка при удалении пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/block_user/<login>', methods=['POST'])
@admin_required
def block_user(login):
    """Блокировка пользователя"""
    try:
        if login == session['user']:
            flash('Нельзя заблокировать самого себя', 'danger')
            return redirect(url_for('admin_panel'))

        success, message = db.block_user(login)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в block_user: {e}")
        flash('Ошибка при блокировке пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/unblock_user/<login>', methods=['POST'])
@admin_required
def unblock_user(login):
    """Разблокировка пользователя"""
    try:
        success, message = db.unblock_user(login)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в unblock_user: {e}")
        flash('Ошибка при разблокировке пользователя', 'danger')
        return redirect(url_for('admin_panel'))


@app.route('/admin/update_user/<login>', methods=['POST'])
@admin_required
def update_user(login):
    """Обновление данных пользователя"""
    try:
        new_data = {
            'email': request.form.get('email', '').strip(),
            'fname': request.form.get('fname', '').strip(),
            'lname': request.form.get('lname', '').strip(),
            'enterprise': request.form.get('enterprise', 'ООО "Верофарм"'),
            'plant': request.form.get('plant', 'Вольгинский'),
            'rights': request.form.get('rights', 'Оператор СиМ'),
            'status': request.form.get('status', 'active')
        }

        new_password = request.form.get('password', '')
        if new_password:
            if len(new_password) < 6:
                flash('Пароль должен содержать минимум 6 символов', 'warning')
                return redirect(url_for('admin_panel'))
            new_data['password'] = new_password

        success, message = db.update_user(login, new_data)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('admin_panel'))
    except Exception as e:
        print(f"Ошибка в update_user: {e}")
        flash('Ошибка при обновлении пользователя', 'danger')
        return redirect(url_for('admin_panel'))


# ============================================================
# МАРШРУТЫ ДЛЯ РАБОТЫ С МАТЕРИАЛАМИ
# ============================================================

@app.route('/materials')
@admin_required
def materials_page():
    """Страница управления материалами"""
    try:
        materials = materials_db.get_all_materials()
        return render_template('materials.html', materials=materials)
    except Exception as e:
        print(f"Ошибка в materials_page: {e}")
        flash('Ошибка загрузки данных материалов', 'danger')
        return redirect(url_for('settings_page'))


@app.route('/materials/update', methods=['POST'])
@admin_required
def update_materials():
    """Обновление данных материалов"""
    try:
        materials_data = request.form.get('materials_data')
        if not materials_data:
            flash('Нет данных для сохранения', 'danger')
            return redirect(url_for('materials_page'))

        materials = json.loads(materials_data)
        errors = []
        success_count = 0

        for material in materials:
            sap_code = material.get('sap_code')
            material_name = material.get('material_name', '').strip()
            palletization = material.get('palletization', 0)
            box_quantity = material.get('box_quantity', 0)

            existing = materials_db.get_material_by_name(material_name)
            if existing and existing['sap_code'] != sap_code:
                errors.append(f"Материал с названием '{material_name}' уже существует (SAP: {existing['sap_code']})")
                continue

            try:
                palletization = int(palletization)
                if palletization < 0:
                    errors.append(f"SAP {sap_code}: Палетизация не может быть отрицательной")
                    continue
            except (ValueError, TypeError):
                errors.append(f"SAP {sap_code}: Палетизация должна быть числом")
                continue

            try:
                box_quantity = int(box_quantity)
                if box_quantity < 0:
                    errors.append(f"SAP {sap_code}: Вложение в короб не может быть отрицательным")
                    continue
            except (ValueError, TypeError):
                errors.append(f"SAP {sap_code}: Вложение в короб должно быть числом")
                continue

            new_data = {
                'material_name': material_name,
                'palletization': palletization,
                'box_quantity': box_quantity
            }
            success, message = materials_db.update_material(sap_code, new_data)
            if success:
                success_count += 1
            else:
                errors.append(message)

        if errors:
            flash(f"Сохранено {success_count} материалов. Ошибки:\n" + "\n".join(errors), 'warning')
        else:
            flash(f"Успешно сохранено {success_count} материалов", 'success')

        return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в update_materials: {e}")
        flash(f'Ошибка при сохранении данных: {str(e)}', 'danger')
        return redirect(url_for('materials_page'))


@app.route('/materials/delete/<int:sap_code>', methods=['POST'])
@admin_required
def delete_material(sap_code):
    """Удаление материала"""
    try:
        success, message = materials_db.delete_material(sap_code)
        flash(message, 'success' if success else 'danger')
        return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в delete_material: {e}")
        flash(f'Ошибка при удалении материала: {str(e)}', 'danger')
        return redirect(url_for('materials_page'))


@app.route('/materials/import', methods=['GET', 'POST'])
@admin_required
def import_materials():
    """Импорт материалов из Excel файла"""
    if request.method == 'GET':
        return render_template('import_materials.html')

    try:
        if 'file' not in request.files:
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_materials'))

        file = request.files['file']

        if file.filename == '':
            flash('Файл не выбран', 'danger')
            return redirect(url_for('import_materials'))

        if not file.filename.lower().endswith(('.xlsx', '.xls')):
            flash('Пожалуйста, загрузите файл в формате Excel (.xlsx или .xls)', 'danger')
            return redirect(url_for('import_materials'))

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            success, message = materials_db.validate_and_import_from_excel(filepath)
            if success:
                flash(message, 'success')
            else:
                flash(message, 'danger')
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)

        return redirect(url_for('materials_page'))
    except Exception as e:
        print(f"Ошибка в import_materials: {e}")
        flash(f'Ошибка при импорте: {str(e)}', 'danger')
        return redirect(url_for('import_materials'))


@app.route('/upload', methods=['POST'])
@labels_access_required
def upload_file():
    """Загрузка и обработка PDF файла"""
    try:
        if 'file' not in request.files:
            return render_template('error.html', error='Файл не выбран')

        file = request.files['file']

        if file.filename == '':
            return render_template('error.html', error='Файл не выбран')

        if not file.filename.lower().endswith('.pdf'):
            return render_template('error.html', error='Пожалуйста, загрузите файл в формате PDF')

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        try:
            data = extract_data_from_pdf(filepath)
            os.remove(filepath)

            if not data or not data['дата_приемки'] and not data['название_поставщика'] and not data['товары']:
                return render_template('error.html',
                                       error='Не удалось извлечь данные из PDF-файла. Проверьте формат файла.')

            # Обогащаем данные о товарах информацией о палетизации
            for product in data['товары']:
                sap_code = product.get('номенклатурный_номер')
                product['sap_code_display'] = sap_code if sap_code else '-'

                if sap_code:
                    try:
                        sap_code_int = int(sap_code)
                        material = materials_db.get_material_by_sap(sap_code_int)
                        if material:
                            product['palletization'] = material.get('palletization', 0)
                            product['material_found'] = True
                        else:
                            product['palletization'] = 0
                            product['material_found'] = False
                    except (ValueError, TypeError):
                        product['palletization'] = 0
                        product['material_found'] = False
                else:
                    product['palletization'] = 0
                    product['material_found'] = False

                # Рассчитываем количество паллет
                quantity = product.get('количество', 0)
                palletization = product.get('palletization', 0)
                product['pallets'] = calculate_pallets(quantity, palletization)

            db.log_action(session['user'], 'UPLOAD_PDF', f"Загружен PDF: {filename}")

            return render_template('result.html', data=data)
        except Exception as e:
            if os.path.exists(filepath):
                os.remove(filepath)
            return render_template('error.html', error=f'Ошибка при обработке файла: {str(e)}')
    except Exception as e:
        print(f"Ошибка в upload_file: {e}")
        return render_template('error.html', error=f'Ошибка при загрузке файла: {str(e)}')


@app.route('/recalculate_pallets', methods=['POST'])
@labels_access_required
def recalculate_pallets():
    """Пересчет количества паллет для всех товаров"""
    try:
        data = request.get_json()
        if not data or 'products' not in data:
            return jsonify({'success': False, 'error': 'Нет данных'})

        products = data.get('products', [])
        recalculated = []

        for product in products:
            quantity = product.get('количество', 0)
            palletization = product.get('palletization', 0)
            pallets = calculate_pallets(quantity, palletization)
            recalculated.append(pallets)

        return jsonify({'success': True, 'pallets': recalculated})
    except Exception as e:
        print(f"Ошибка в recalculate_pallets: {e}")
        return jsonify({'success': False, 'error': str(e)})


@app.route('/generate_labels', methods=['POST'])
@labels_access_required
def generate_labels():
    """Генерация палетных этикеток"""
    try:
        products_json = request.form.get('products')
        order_number = request.form.get('order_number')

        if not products_json or not order_number:
            return render_template('error.html', error='Нет данных для формирования этикеток')

        products_json = products_json.replace("'", '"')

        try:
            products = json.loads(products_json)
        except json.JSONDecodeError as e:
            print(f"Ошибка парсинга JSON: {e}")
            return render_template('error.html', error=f'Ошибка парсинга данных: {str(e)}')

        if not products:
            return render_template('error.html', error='Нет товаров для формирования этикеток')

        settings = load_settings()
        save_path = settings.get('save_path')
        sop_code = settings.get('sop_code', 'RUPD.VLG.05.04.C01')

        filepath = create_pallet_labels_file(products, order_number, save_path, sop_code)
        open_in_word(filepath)

        db.log_action(session['user'], 'GENERATE_LABELS',
                      f"Сгенерировано этикеток для заказа {order_number}")

        # Подсчитываем общее количество сгенерированных этикеток
        total_pallets = sum([p.get('pallets', 1) for p in products])

        return render_template('success.html',
                               filepath=filepath,
                               product_count=len(products),
                               order_number=order_number,
                               total_pallets=total_pallets)
    except Exception as e:
        print(f"Ошибка в generate_labels: {e}")
        return render_template('error.html', error=f'Ошибка при формировании этикеток: {str(e)}')


@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings_page():
    """Страница настроек - доступна всем пользователям"""
    try:
        settings = load_settings()
        drives = get_available_drives()

        current_path = settings.get('save_path',
                                    os.path.join(os.path.expanduser('~'), 'Documents', 'Палетные этикетки'))

        path_param = request.args.get('path')
        if path_param and os.path.exists(path_param):
            current_path = path_param

        path_parts = []
        if current_path:
            normalized_path = os.path.normpath(current_path)
            parts = normalized_path.split(os.sep)
            temp_path = ""
            for part in parts:
                if part:
                    if temp_path:
                        temp_path = os.path.join(temp_path, part)
                    else:
                        if sys.platform == 'win32' and ':' in part:
                            temp_path = part + os.sep
                        else:
                            temp_path = os.sep + part
                    path_parts.append({
                        'name': part,
                        'path': temp_path
                    })

        subdirs = get_subdirectories(current_path)

        if request.method == 'POST':
            selected_path = request.form.get('selected_path')
            if selected_path:
                settings['save_path'] = selected_path
                save_settings(settings)
                os.makedirs(selected_path, exist_ok=True)
                flash('Настройки сохранены', 'success')
                return redirect(url_for('settings_page'))

            sop_code = request.form.get('sop_code')
            if sop_code:
                settings['sop_code'] = sop_code.strip()
                save_settings(settings)
                flash('Номер СОП успешно обновлен', 'success')
                return redirect(url_for('settings_page'))

            path_component = request.form.get('path_component')
            if path_component:
                new_path = path_component
                settings['save_path'] = new_path
                save_settings(settings)
                os.makedirs(new_path, exist_ok=True)
                flash('Настройки сохранены', 'success')
                return redirect(url_for('settings_page'))

        material_count = materials_db.get_materials_count()

        return render_template('settings.html',
                               settings=settings,
                               drives=drives,
                               current_path=current_path,
                               path_parts=path_parts,
                               subdirs=subdirs,
                               os_sep=os.sep,
                               is_admin=(session.get('user_rights') == 'Администратор'),
                               material_count=material_count)
    except Exception as e:
        print(f"Ошибка в settings_page: {e}")
        import traceback
        traceback.print_exc()
        flash('Ошибка загрузки настроек', 'danger')
        return redirect(url_for('index'))


@app.route('/exit')
def exit_app():
    """Страница выхода"""
    return render_template('exit.html')


if __name__ == '__main__':
    load_settings()
    app.run(debug=True, host='0.0.0.0', port=5000)
# app/services/label_generator.py
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from app.utils.helpers import clean_filename, get_unique_filename


def get_name_length_category(name_len, font_settings):
    """Определяет категорию длины названия"""
    short_threshold = font_settings.get('short_name_threshold', 25)
    medium_threshold = font_settings.get('medium_name_threshold', 40)

    if name_len < short_threshold:
        return 'short'
    elif name_len < medium_threshold:
        return 'medium'
    else:
        return 'long'


def get_font_and_spacing_settings(product_name, font_settings):
    """Определяет размеры шрифтов и межстрочный интервал"""
    name_len = len(product_name)
    category = get_name_length_category(name_len, font_settings)

    if category == 'short':
        return {
            'header': font_settings.get('header_size', 12),
            'product': font_settings.get('product_size_short', 80),
            'supplier': font_settings.get('supplier_size_short', 42),
            'analytic': font_settings.get('analytic_size', 117),
            'code': font_settings.get('code_size', 14),
            'line_spacing': font_settings.get('line_spacing_short', 1.5),
            'category': 'short'
        }
    elif category == 'medium':
        return {
            'header': font_settings.get('header_size', 12),
            'product': font_settings.get('product_size_medium', 72),
            'supplier': font_settings.get('supplier_size_medium', 42),
            'analytic': font_settings.get('analytic_size', 117),
            'code': font_settings.get('code_size', 14),
            'line_spacing': font_settings.get('line_spacing_medium', 1.0),
            'category': 'medium'
        }
    else:
        return {
            'header': font_settings.get('header_size', 12),
            'product': font_settings.get('product_size_long', 64),
            'supplier': font_settings.get('supplier_size_long', 36),
            'analytic': font_settings.get('analytic_size', 117),
            'code': font_settings.get('code_size', 14),
            'line_spacing': font_settings.get('line_spacing_long', 0.8),
            'category': 'long'
        }


def create_pallet_labels_file(all_products, order_number, supplier_name, order_date,
                              selected_indices, save_path, sop_code, font_settings):
    """Создает файл с палетными этикетками только для выбранных товаров"""

    try:
        os.makedirs(save_path, exist_ok=True)
    except Exception:
        save_path = tempfile.gettempdir()
        os.makedirs(save_path, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.15)
        section.bottom_margin = Inches(0.15)
        section.left_margin = Inches(0.15)
        section.right_margin = Inches(0.15)

    total_labels = 0
    first_page = True

    for idx, product in enumerate(all_products):
        if idx not in selected_indices:
            continue

        pallets_count = product.get('pallets', 1)
        if pallets_count < 1:
            pallets_count = 1

        product_name = product.get('наименование', 'Не указано')
        supplier_series = product.get('партия_поставщика', 'Не указана')
        analytic_list = product.get('аналитический_лист', 'Не указан')

        font_settings_dict = get_font_and_spacing_settings(product_name, font_settings)
        line_spacing = font_settings_dict.get('line_spacing', 1.0)

        for pallet_num in range(pallets_count):
            if not first_page:
                doc.add_page_break()
                if doc.sections:
                    section = doc.sections[-1]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11.69)
                    section.page_height = Inches(8.27)
                    section.top_margin = Inches(0.15)
                    section.bottom_margin = Inches(0.15)
                    section.left_margin = Inches(0.15)
                    section.right_margin = Inches(0.15)
            else:
                first_page = False

            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(10.8)

            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].clear()

            # ООО «Верофарм»
            p1 = cell.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(2)
            p1.paragraph_format.space_after = Pt(2)
            p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p1.paragraph_format.line_spacing = line_spacing
            run1 = p1.add_run('ООО «Верофарм»')
            run1.italic = True
            run1.bold = True
            run1.font.size = Pt(font_settings_dict['header'])
            run1.font.name = 'Calibri'

            # Название товара
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(5)
            p2.paragraph_format.space_after = Pt(5)
            p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p2.paragraph_format.line_spacing = line_spacing
            run2 = p2.add_run(product_name)
            run2.italic = True
            run2.bold = True
            run2.font.size = Pt(font_settings_dict['product'])
            run2.font.name = 'Calibri'

            # Партия поставщика
            supplier_text = f"п. {supplier_series}" if supplier_series and supplier_series != 'Не указана' else supplier_series
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(4)
            p3.paragraph_format.space_after = Pt(4)
            p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p3.paragraph_format.line_spacing = line_spacing
            run3 = p3.add_run(supplier_text)
            run3.italic = True
            run3.bold = False
            run3.font.size = Pt(font_settings_dict['supplier'])
            run3.font.name = 'Calibri'

            # Аналитический лист
            p4 = cell.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(4)
            p4.paragraph_format.space_after = Pt(4)
            p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p4.paragraph_format.line_spacing = line_spacing
            run4 = p4.add_run(f'АЛ {analytic_list}')
            run4.italic = True
            run4.bold = True
            run4.font.size = Pt(font_settings_dict['analytic'])
            run4.font.name = 'Calibri'

            # Номер СОП
            p5 = cell.add_paragraph()
            p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p5.paragraph_format.space_before = Pt(8)
            p5.paragraph_format.space_after = Pt(2)
            p5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p5.paragraph_format.line_spacing = line_spacing
            run5 = p5.add_run(sop_code)
            run5.italic = False
            run5.bold = True
            run5.font.size = Pt(font_settings_dict['code'])
            run5.font.name = 'Calibri'

            total_labels += 1

    clean_supplier = clean_filename(supplier_name) if supplier_name else 'Неизвестный_поставщик'

    date_str = ''
    if order_date:
        import re
        date_match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', order_date)
        if date_match:
            date_str = f"{date_match.group(1)}{date_match.group(2)}{date_match.group(3)}"
        else:
            from datetime import datetime
            date_str = datetime.now().strftime("%d%m%Y")
    else:
        from datetime import datetime
        date_str = datetime.now().strftime("%d%m%Y")

    base_name = f"{clean_supplier}_{order_number}_{date_str}"
    filepath = get_unique_filename(save_path, base_name, '.docx')

    try:
        doc.save(filepath)
    except Exception:
        temp_dir = tempfile.gettempdir()
        filepath = get_unique_filename(temp_dir, base_name, '.docx')
        doc.save(filepath)

    return filepath, total_labels